# -*- coding: utf-8 -*-
"""Result v1.4 generator + finalizer + quality loop.

Rules:
- punctuation: body-only Chinese quotes, headings/references unchanged
- run-level CJK/Latin split fonts
- heading rules (H1/H2/H3)
- table rules (Table Grid + table content style, 85%, center, content-aware columns)
- caption style + keepNext
- references: template numbering, hanging 0.74cm, 12pt
- TOC single field + settings updateFields
"""
import io
import json
import os
import re
import shutil
import sys
import zipfile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from lxml import etree


SONG = "\u5b8b\u4f53"
TNR = "Times New Roman"
WEI = "\u5fae\u8f6f\u96c5\u9ed1"


CJK_PUNCT = set(
    "\u201c\u201d\u2018\u2019\uff0c\u3002\u3001\uff1b\uff1a\uff1f\uff01"
    "\uff08\uff09\u3010\u3011\u300a\u300b\u3008\u3009\u2014\u2026\uff5e"
)


def _cjk(ch):
    return ch in CJK_PUNCT or ('\u4e00' <= ch <= '\u9fff')


def split_text(text):
    segs = []
    buf = ""
    mode = None
    for ch in text:
        m = _cjk(ch)
        if mode is None:
            mode = m
        if m != mode:
            if buf:
                segs.append((buf, mode))
            buf = ch
            mode = m
        else:
            buf += ch
    if buf:
        segs.append((buf, mode))
    return segs


def _set_font(run, cjk, size, bold, latin, is_cjk=True):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = cjk if is_cjk else latin
    rPr = run._r.get_or_add_rPr()
    rf = rPr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rPr.append(rf)
    rf.set(qn("w:ascii"), cjk if is_cjk else latin)
    rf.set(qn("w:hAnsi"), cjk if is_cjk else latin)
    rf.set(qn("w:eastAsia"), cjk)
    if is_cjk:
        rf.set(qn("w:hint"), "eastAsia")


def rebuild_runs(paragraph, cjk, size, bold, latin):
    text = paragraph.text
    for r in list(paragraph.runs):
        r._r.getparent().remove(r._r)
    for seg, is_cjk in split_text(text):
        run = paragraph.add_run(seg)
        t = run._r.find(qn("w:t"))
        if t is not None:
            t.set(qn("xml:space"), "preserve")
        _set_font(run, cjk if is_cjk else SONG, size, bold, latin, is_cjk)




def _set_paragraph_text(p, text):
    el = p._p
    for child in list(el):
        if child.tag in (qn("w:r"), qn("w:hyperlink")):
            el.remove(child)
    run = p.add_run(text)
    t = run._r.find(qn("w:t"))
    if t is not None:
        t.set(qn("xml:space"), "preserve")
    return run


REAL_NETWORK_REPLACEMENTS = [
    ("\u957f\u6c99\u57ce\u5e02\u8f68\u9053\u4ea4\u901a", "\u67d0\u57ce\u5e02\u8f68\u9053\u4ea4\u901a\u7ebf\u8def"),
    ("\u957f\u6c99\u5730\u94c1\u7f51\u7edc", "\u67d0\u8f68\u9053\u4ea4\u901a\u7ebf\u8def\u7f51\u7edc"),
    ("\u957f\u6c99\u8f68\u9053\u4ea4\u901a\u7f51\u7edc", "\u67d0\u8f68\u9053\u4ea4\u901a\u7ebf\u8def\u7f51\u7edc"),
    ("\u957f\u6c99\u5730\u94c1", "\u67d0\u8f68\u9053\u4ea4\u901a\u7ebf\u8def"),
    ("\u957f\u6c99\u8f68\u9053\u4ea4\u901a", "\u67d0\u8f68\u9053\u4ea4\u901a\u7ebf\u8def"),
]


def neutralize_content(doc):
    def fix(p):
        text = p.text
        newtext = text
        for a, b in REAL_NETWORK_REPLACEMENTS:
            newtext = newtext.replace(a, b)
        if newtext != text:
            _set_paragraph_text(p, newtext)
    started = False
    for p in doc.paragraphs:
        style = p.style.name if p.style else ""
        t = p.text.strip()
        if not started and style.startswith("Heading") and re.match(r"^1\s", t):
            started = True
            continue
        if not started:
            continue
        if t == "\u53c2\u8003\u6587\u732e":
            break
        fix(p)
    for ti in range(2, len(doc.tables)):
        for row in doc.tables[ti].rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    fix(p)


def body_punctuation(doc):
    def fix(p):
        text = p.text
        new = text
        if '"' in new or "'" in new:
            new = re.sub(r'"([^"\n]+)"', "\u201c\\1\u201d", new)
            new = re.sub(r"'([^'\n]+)'", "\u2018\\1\u2019", new)
        full = {",": "\uff0c", ";": "\uff1b", ":": "\uff1a", "?": "\uff1f", "!": "\uff01",
                "(": "\uff08", ")": "\uff09"}
        for a, b in full.items():
            new = new.replace(a, b)
        if new != text:
            _set_paragraph_text(p, new)
    started = False
    for p in doc.paragraphs:
        style = p.style.name if p.style else ""
        t = p.text.strip()
        if not started and style.startswith("Heading") and re.match(r"^1\s", t):
            started = True
            continue
        if not started:
            continue
        if t == "\u53c2\u8003\u6587\u732e":
            break
        if style.startswith("Heading"):
            continue
        if re.match(r"^\[?\d", t) or re.match(r"^\u8868\d", t):
            continue
        fix(p)
    for ti in range(2, len(doc.tables)):
        for row in doc.tables[ti].rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        fix(p)

def heading_rules(doc):
    for p in doc.paragraphs:
        style = p.style.name if p.style else ""
        if not style.startswith("Heading"):
            continue
        level = int(style.split()[-1])
        if level == 1:
            p.alignment = 1  # center
            pPr = p._p.get_or_add_pPr()
            sp = pPr.find(qn("w:spacing"))
            if sp is None:
                sp = OxmlElement("w:spacing")
                pPr.append(sp)
            sp.set(qn("w:beforeLines"), "100")
            sp.set(qn("w:afterLines"), "100")
            sp.set(qn("w:before"), "240")
            sp.set(qn("w:after"), "240")
        elif level == 2:
            p.alignment = 3  # justify




def list_number_spacing(doc):
    started = False
    for p in doc.paragraphs:
        style = p.style.name if p.style else ""
        t = p.text.strip()
        if not started and style.startswith("Heading") and re.match(r"^1\s", t):
            started = True
            continue
        if not started:
            continue
        if t == "\u53c2\u8003\u6587\u732e":
            break
        if style.startswith("Heading") or not t:
            continue
        newtext = re.sub(r"^(\d+[.\u3001)\uff0e])\s+", r"\1", t)
        newtext = re.sub(r"^[\uff08(](\d+)[\uff09)]\s+", r"(\1)", newtext)
        if newtext != t:
            _set_paragraph_text(p, newtext)


def body_fonts(doc):
    started = False
    for p in doc.paragraphs:
        style = p.style.name if p.style else ""
        t = p.text.strip()
        if not started and style.startswith("Heading") and re.match(r"^1\s", t):
            started = True
            continue
        if not started:
            continue
        if t == "\u53c2\u8003\u6587\u732e":
            break
        if style.startswith("Heading") or not t:
            continue
        if re.match(r"^\[?\d", t) or re.match(r"^\u8868\d", t):
            continue
        rebuild_runs(p, SONG, 12, False, TNR)

def _col_weights(table):
    ncols = len(table.columns)
    weights = [1.0] * ncols
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells[:ncols]):
            units = sum(2 if _cjk(c) else 1 for c in cell.text[:80])
            weights[ci] = max(weights[ci], units + 6)
    return weights




def table_rules(doc):
    for ti in (2, 3, 4, 5):
        if ti >= len(doc.tables):
            continue
        t = doc.tables[ti]
        tblPr = t._tbl.tblPr
        ts = tblPr.find(qn("w:tblStyle"))
        if ts is None:
            ts = OxmlElement("w:tblStyle")
            tblPr.insert(0, ts)
        ts.set(qn("w:val"), "18")
        tw = tblPr.find(qn("w:tblW"))
        if tw is None:
            tw = OxmlElement("w:tblW")
            tblPr.insert(0, tw)
        tw.set(qn("w:w"), "4998")
        tw.set(qn("w:type"), "pct")
        jc = tblPr.find(qn("w:jc"))
        if jc is not None:
            tblPr.remove(jc)
        weights = _col_weights(t)
        total_w = 9068
        s = sum(weights)
        col_widths = []
        grid = t._tbl.find(qn("w:tblGrid"))
        if grid is not None:
            for ci, c in enumerate(grid.findall(qn("w:gridCol"))):
                if ci < len(weights):
                    c.set(qn("w:w"), str(int(round(total_w * weights[ci] / s))))
                col_widths.append(int(c.get(qn("w:w"), "1000")))
        seen = set()
        for row in t.rows:
            for ci, c in enumerate(row.cells):
                if c._tc in seen:
                    continue
                seen.add(c._tc)
                tcPr = c._tc.find(qn("w:tcPr"))
                if tcPr is None:
                    tcPr = OxmlElement("w:tcPr")
                    c._tc.insert(0, tcPr)
                shd = tcPr.find(qn("w:shd"))
                if shd is None:
                    shd = OxmlElement("w:shd")
                    tcPr.append(shd)
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "FFFFFF")
                if ci < len(col_widths):
                    tw = tcPr.find(qn("w:tcW"))
                    if tw is None:
                        tw = OxmlElement("w:tcW")
                        tcPr.insert(0, tw)
                    tw.set(qn("w:w"), str(col_widths[ci]))
                    tw.set(qn("w:type"), "dxa")
                va = tcPr.find(qn("w:vAlign"))
                if va is None:
                    va = OxmlElement("w:vAlign")
                    tcPr.append(va)
                va.set(qn("w:val"), "center")
                tb = tcPr.find(qn("w:tcBorders"))
                if tb is not None:
                    tcPr.remove(tb)
                tb = OxmlElement("w:tcBorders")
                for edge in ("top", "left", "bottom", "right"):
                    e = OxmlElement("w:%s" % edge)
                    e.set(qn("w:val"), "single")
                    e.set(qn("w:color"), "000000")
                    e.set(qn("w:sz"), "4")
                    e.set(qn("w:space"), "0")
                    tb.append(e)
                vael = tcPr.find(qn("w:vAlign"))
                if vael is not None:
                    vael.addprevious(tb)
                else:
                    tcPr.append(tb)
                for p in c.paragraphs:
                    pPr = p._p.get_or_add_pPr()
                    ps = pPr.find(qn("w:pStyle"))
                    if ps is None:
                        ps = OxmlElement("w:pStyle")
                        pPr.insert(0, ps)
                    ps.set(qn("w:val"), "33")
                    jc2 = pPr.find(qn("w:jc"))
                    if jc2 is None:
                        jc2 = OxmlElement("w:jc")
                        pPr.append(jc2)
                    jc2.set(qn("w:val"), "center")
                    rebuild_runs(p, SONG, 10.5, False, TNR)



def caption_rules(doc):
    biaozhu = None
    try:
        biaozhu = doc.styles["\u8868\u6ce8"]
    except Exception:
        pass
    for p in doc.paragraphs:
        style = p.style.name if p.style else ""
        t = p.text.strip()
        if style == "\u8868\u6ce8" or (t.startswith("\u8868") and len(t) > 1 and t[1].isdigit()):
            if biaozhu is not None:
                try:
                    p.style = biaozhu
                except Exception:
                    pass
            pPr = p._p.get_or_add_pPr()
            kn = pPr.find(qn("w:keepNext"))
            if kn is None:
                kn = OxmlElement("w:keepNext")
                pPr.insert(0, kn)
            newtext = re.sub(r"^\u8868(\d+)[-.]\s*(\d+)\s*", lambda m: "\u8868%s.%s  " % (m.group(1), m.group(2)), t)
            _set_paragraph_text(p, newtext)
            rebuild_runs(p, SONG, 10.5, False, TNR)

def reference_rules(doc):
    in_ref = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if t == "\u53c2\u8003\u6587\u732e":
            in_ref = True
            continue
        if not in_ref or not t:
            continue
        pPr = p._p.get_or_add_pPr()
        numPr = pPr.find(qn("w:numPr"))
        if numPr is None:
            numPr = OxmlElement("w:numPr")
            pPr.insert(0, numPr)
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        numId = OxmlElement("w:numId")
        numId.set(qn("w:val"), "1")
        numPr.append(ilvl)
        numPr.append(numId)
        ind = pPr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind")
            pPr.append(ind)
        ind.set(qn("w:left"), "420")
        ind.set(qn("w:hanging"), "420")
        ind.set(qn("w:firstLine"), "0")
        newtext = re.sub(r"^\[\d+\]\s*", "", p.text)
        _set_paragraph_text(p, newtext)
        rebuild_runs(p, SONG, 12, False, TNR)


def rebuild_toc(doc):
    body = doc.element.body
    children = list(body.iterchildren())
    start_idx = end_idx = None
    for i, el in enumerate(children):
        if el.tag != qn("w:p"):
            continue
        p = Paragraph(el, doc)
        t = p.text.strip()
        if start_idx is None and t.replace(" ", "") == "\u76ee\u5f55":
            start_idx = i
            continue
        if start_idx is not None and end_idx is None:
            if (p.style.name if p.style else "").startswith("Heading") and re.match(r"^1\s", t):
                end_idx = i
                break
    if start_idx is not None and end_idx is not None and end_idx > start_idx + 1:
        kept_sect = []
        for el in children[start_idx + 1:end_idx]:
            if el.tag != qn("w:p"):
                continue
            pPr = el.find(qn("w:pPr"))
            if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
                kept_sect.append(el)
        for el in children[start_idx + 1:end_idx]:
            if el not in kept_sect:
                body.remove(el)
    anchor = children[start_idx]
    new_p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    ps = OxmlElement("w:pStyle")
    ps.set(qn("w:val"), "TOC 1")
    pPr.append(ps)
    new_p.append(pPr)

    def fc(tp):
        r = OxmlElement("w:r")
        e = OxmlElement("w:fldChar")
        e.set(qn("w:fldCharType"), tp)
        r.append(e)
        return r

    def ins(x):
        r = OxmlElement("w:r")
        e = OxmlElement("w:instrText")
        e.set(qn("xml:space"), "preserve")
        e.text = x
        r.append(e)
        return r

    new_p.append(fc("begin"))
    new_p.append(ins(' TOC \\o "1-3" \\h \\u '))
    new_p.append(fc("separate"))
    new_p.append(fc("end"))
    anchor.addnext(new_p)



def _toc_style_id(level):
    return {1: "13", 2: "15", 3: "7"}.get(level, "13")


def _toc_pPr(level):
    pPr = OxmlElement("w:pPr")
    ps = OxmlElement("w:pStyle")
    ps.set(qn("w:val"), _toc_style_id(level))
    pPr.append(ps)
    tabs = OxmlElement("w:tabs")
    t1 = OxmlElement("w:tab")
    t1.set(qn("w:val"), "right")
    t1.set(qn("w:leader"), "dot")
    t1.set(qn("w:pos"), "9355")
    tabs.append(t1)
    if level == 1:
        t2 = OxmlElement("w:tab")
        t2.set(qn("w:val"), "clear")
        t2.set(qn("w:pos"), "8777")
        tabs.append(t2)
    pPr.append(tabs)
    return pPr


def _fld_run(ftype):
    r = OxmlElement("w:r")
    fc = OxmlElement("w:fldChar")
    fc.set(qn("w:fldCharType"), ftype)
    r.append(fc)
    return r


def _instr_run(text):
    r = OxmlElement("w:r")
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = text
    r.append(it)
    return r


def _text_run(text):
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def _tab_run():
    r = OxmlElement("w:r")
    t = OxmlElement("w:tab")
    r.append(t)
    return r


def _collect_headings(doc):
    out = []
    for p in doc.paragraphs:
        style = p.style.name if p.style else ""
        if not style.startswith("Heading"):
            continue
        try:
            level = int(style.split()[-1])
        except Exception:
            continue
        if level not in (1, 2, 3):
            continue
        text = p.text.strip()
        if text:
            out.append((level, text, p))
    return out


def _norm_key(s):
    return "".join(s.split()).replace("\u3000", "")


def _pdf_heading_pages(pdf_path, headings):
    try:
        import pdfplumber
    except Exception:
        return {}
    pages = []
    with pdfplumber.open(pdf_path) as pdf:
        for pg in pdf.pages:
            txt = (pg.extract_text() or "").strip()
            lines = [ln.strip() for ln in txt.splitlines() if ln.strip()]
            disp = None
            if lines:
                last = lines[-1].replace(" ", "").replace("\u3000", "")
                if last.isdigit():
                    disp = int(last)
            pages.append((txt, disp))
    result = {}
    start = 0
    for level, text, p in headings:
        key = _norm_key(text)
        for i in range(start, len(pages)):
            txt, disp = pages[i]
            if disp is None:
                continue
            if key in _norm_key(txt):
                result[key] = disp
                start = i
                break
    return result


def _add_toc_bookmark(paragraph, bid):
    p = paragraph._p
    pPr = p.find(qn("w:pPr"))
    bs = OxmlElement("w:bookmarkStart")
    bs.set(qn("w:id"), str(5000 + bid))
    bs.set(qn("w:name"), "_Toc%d" % bid)
    be = OxmlElement("w:bookmarkEnd")
    be.set(qn("w:id"), str(5000 + bid))
    if pPr is not None:
        pPr.addnext(bs)
    else:
        p.insert(0, bs)
    p.append(be)


def toc_cache(docx_path, page_map):
    doc = Document(docx_path)
    body = doc.element.body
    toc_p = None
    for el in body.iterchildren():
        if el.tag != qn("w:p"):
            continue
        xml = etree.tostring(el, encoding="unicode")
        if "TOC \\o" in xml:
            toc_p = el
            break
    if toc_p is None:
        return False
    headings = _collect_headings(doc)
    if not headings:
        return False
    for r in list(toc_p.findall(qn("w:r"))):
        fc = r.find(qn("w:fldChar"))
        if fc is not None and fc.get(qn("w:fldCharType")) == "end":
            toc_p.remove(r)
    sdt = OxmlElement("w:sdt")
    sdtPr = OxmlElement("w:sdtPr")
    dpo = OxmlElement("w:docPartObj")
    dpg = OxmlElement("w:docPartGallery")
    dpg.set(qn("w:val"), "Table of Contents")
    dpu = OxmlElement("w:docPartUnique")
    dpo.append(dpg)
    dpo.append(dpu)
    sdtPr.append(dpo)
    sdtContent = OxmlElement("w:sdtContent")
    sdt.append(sdtPr)
    sdt.append(sdtContent)
    old_pPr = toc_p.find(qn("w:pPr"))
    if old_pPr is not None:
        toc_p.remove(old_pPr)
    toc_p.insert(0, _toc_pPr(headings[0][0]))
    toc_p.addprevious(sdt)
    sdtContent.append(toc_p)
    prev = toc_p
    bid = 1001
    for idx, (level, text, p) in enumerate(headings):
        key = _norm_key(text)
        page = page_map.get(key, 1)
        runs = [_fld_run("begin"),
                _instr_run(' HYPERLINK \\l "_Toc%d" ' % bid),
                _fld_run("separate"),
                _text_run(text),
                _tab_run(),
                _fld_run("begin"),
                _instr_run(' PAGEREF _Toc%d \\h ' % bid),
                _fld_run("separate"),
                _text_run(str(page)),
                _fld_run("end"),
                _fld_run("end")]
        if idx == 0:
            for r in runs:
                toc_p.append(r)
            _add_toc_bookmark(p, bid)
        else:
            new_p = OxmlElement("w:p")
            new_p.append(_toc_pPr(level))
            for r in runs:
                new_p.append(r)
            sdtContent.append(new_p)
            prev = new_p
            _add_toc_bookmark(p, bid)
        bid += 1
    prev.append(_fld_run("end"))
    doc.save(docx_path)
    return True



def ensure_toc_styles(docx_path):
    tmp = docx_path + ".tmpstyles"
    with zipfile.ZipFile(docx_path) as zin:
        with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
            seen = set()
            for item in zin.infolist():
                if item.filename in seen:
                    continue
                seen.add(item.filename)
                data = zin.read(item.filename)
                if item.filename == "word/styles.xml":
                    root = etree.fromstring(data)
                    for st in root.iter(qn("w:style")):
                        if st.get(qn("w:styleId")) == "13":
                            rpr = st.find(qn("w:rPr"))
                            if rpr is None:
                                rpr = OxmlElement("w:rPr")
                                st.append(rpr)
                            rf = rpr.find(qn("w:rFonts"))
                            if rf is None:
                                rf = OxmlElement("w:rFonts")
                                rpr.insert(0, rf)
                            rf.set(qn("w:ascii"), "Times New Roman")
                            rf.set(qn("w:hAnsi"), "Times New Roman")
                            rf.set(qn("w:eastAsia"), SONG)
                            rf.set(qn("w:cs"), "Times New Roman")
                    data = etree.tostring(root, encoding="utf-8", xml_declaration=True, standalone=True)
                zout.writestr(item, data)
    shutil.move(tmp, docx_path)


def ensure_settings(docx_path):
    tmp = docx_path + ".tmp"
    with zipfile.ZipFile(docx_path) as zin:
        with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
            seen = set()
            for item in zin.infolist():
                if item.filename in seen:
                    continue
                seen.add(item.filename)
                data = zin.read(item.filename)
                if item.filename == "word/settings.xml":
                    s = data.decode("utf-8")
                    if "updateFields" not in s:
                        s = s.replace("</w:settings>",
                                      '<w:updateFields w:val="true"/></w:settings>')
                    data = s.encode("utf-8")
                zout.writestr(item, data)
    shutil.move(tmp, docx_path)


def dedup_numbering(docx_path):
    tmp = docx_path + ".tmpz"
    with zipfile.ZipFile(docx_path) as zin:
        seen = set()
        with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename in seen:
                    continue
                seen.add(item.filename)
                zout.writestr(item, zin.read(item.filename))
    shutil.move(tmp, docx_path)


def base_generate(info, stu_dir, direction, out_docx, golden, ai_docx, rrb):
    stu = info.get("\u59d3\u540d", "")
    topic = info.get("\u8bfe\u9898\u540d\u79f0", "")
    rrb.reconstruct(stu, direction, "02", info, ai_docx, out_docx, neutralize=True)
    d = Document(out_docx)
    body = d.element.body
    bs = [el for el in body if el.tag == qn("w:sectPr")]
    if len(bs) > 1:
        body.remove(bs[0])
    # heading remap + spacing text
    for p in d.paragraphs:
        style = p.style.name if p.style else ""
        if "Heading" not in style:
            continue
        t = p.text.strip()
        m = re.match(r"^(\d+(?:\.\d+)*)\s*(.*)$", t)
        if not m:
            continue
        num, title = m.group(1), m.group(2).strip()
        depth = 1
        if re.match(r"^\d+\.\d+\.\d+", num):
            depth = 3
        elif re.match(r"^\d+\.\d+", num):
            depth = 2
        if depth == 1 and len(title) == 2 and all(_cjk(c) for c in title):
            title = title[0] + "  " + title[1]
        newtext = num + "  " + title
        try:
            p.style = d.styles["Heading %d" % depth]
        except Exception:
            pass
        _set_paragraph_text(p, newtext)
    neutralize_content(d)
    # clear source formatting in content area
    started = False
    in_ref = False
    for p in d.paragraphs:
        style = p.style.name if p.style else ""
        t = p.text.strip()
        if not started and style.startswith("Heading") and re.match(r"^1\s", t):
            started = True
        if t == "\u53c2\u8003\u6587\u732e":
            in_ref = True
        if not started:
            continue
        pPr = p._p.find(qn("w:pPr"))
        if pPr is not None:
            for child in list(pPr):
                if child.tag != qn("w:pStyle"):
                    pPr.remove(child)
        for r in p.runs:
            rPr = r._r.find(qn("w:rPr"))
            if rPr is not None:
                r._r.remove(rPr)
    d.save(out_docx)
    with zipfile.ZipFile(golden) as z:
        numxml = z.read("word/numbering.xml")
    tmp = out_docx + ".tmp"
    with zipfile.ZipFile(out_docx) as zin:
        seen = set()
        with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename in seen:
                    continue
                seen.add(item.filename)
                data = zin.read(item.filename)
                if item.filename == "[Content_Types].xml":
                    s = data.decode("utf-8")
                    if "numbering" not in s:
                        s = s.replace("</Types>",
                                      '<Override PartName="/word/numbering.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml"/></Types>')
                    data = s.encode("utf-8")
                elif item.filename == "word/_rels/document.xml.rels":
                    s = data.decode("utf-8")
                    if "numbering.xml" not in s:
                        s = s.replace("</Relationships>",
                                      '<Relationship Id="rIdNumbering99" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering" Target="numbering.xml"/></Relationships>')
                    data = s.encode("utf-8")
                zout.writestr(item, data)
            zout.writestr("word/numbering.xml", numxml)
    shutil.move(tmp, out_docx)
    dedup_numbering(out_docx)



def _template_page_footer_parts(golden):
    if not golden or not os.path.isfile(golden):
        return None
    names = set()
    try:
        with zipfile.ZipFile(golden) as z:
            for n in z.namelist():
                if re.match(r"^word/footer\d+\.xml$", n):
                    data = z.read(n).decode("utf-8")
                    if re.search(r"<w:instrText[^>]*>\s*PAGE\s*</w:instrText>", data):
                        names.add(n)
    except Exception:
        return None
    return names


def ensure_page_fields(doc, golden=None):
    template_pages = _template_page_footer_parts(golden)
    for section in doc.sections:
        footer = section.footer
        if template_pages is not None:
            if footer.is_linked_to_previous:
                continue
            try:
                part_name = str(footer.part.partname).lstrip("/")
            except Exception:
                continue
            if part_name not in template_pages:
                continue
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        if "PAGE" in (p._p.xml or ""):
            continue
        p.alignment = 1
        run = p.add_run()
        f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
        ins = OxmlElement("w:instrText"); ins.set(qn("xml:space"), "preserve"); ins.text = " PAGE "
        f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end")
        run._r.append(f1); run._r.append(ins); run._r.append(f2)
        run.font.size = Pt(10.5); run.font.name = "Times New Roman"


def finalize(out_docx, golden=None):
    d = Document(out_docx)
    heading_rules(d)
    body_punctuation(d)
    list_number_spacing(d)
    body_fonts(d)
    table_rules(d)
    caption_rules(d)
    reference_rules(d)
    rebuild_toc(d)
    ensure_page_fields(d, golden)
    d.save(out_docx)
    ensure_settings(out_docx)
    ensure_toc_styles(out_docx)


def main():
    info_path = sys.argv[1]
    direction = sys.argv[2]
    out_docx = sys.argv[3]
    golden = sys.argv[4]
    with io.open(info_path, encoding="utf-8") as f:
        info = json.load(f)
    stu_dir = os.path.dirname(info_path)
    stu = info.get("\u59d3\u540d", "")
    ai_docx = os.path.join(stu_dir, "\u6210\u679c\u521d\u7a3f.docx")
    ws = os.path.dirname(os.path.dirname(golden))
    sys.path.insert(0, os.path.join(ws, "00_\u7cfb\u7edf\u914d\u7f6e", "\u6a21\u5757", "v03"))
    import result_reference_builder as rrb
    base_generate(info, stu_dir, direction, out_docx, golden, ai_docx, rrb)
    finalize(out_docx, golden)
    pdf_out = sys.argv[5] if len(sys.argv) > 5 else None
    final_pdf = None
    try:
        v03_dir = os.path.join(ws, "00_\u7cfb\u7edf\u914d\u7f6e", "\u6a21\u5757", "v03")
        sys.path.insert(0, os.path.join(os.path.dirname(v03_dir), "v06"))
        import render_docx as rdx
        import tempfile
        tmp_dir = tempfile.mkdtemp(prefix="result_toc_")
        tmp_src = os.path.join(tmp_dir, "src.docx")
        shutil.copy2(out_docx, tmp_src)
        tmp_pdf = os.path.join(tmp_dir, "preview.pdf")
        if rdx.render_to_pdf(tmp_src, tmp_pdf):
            d = Document(out_docx)
            pmap = _pdf_heading_pages(tmp_pdf, _collect_headings(d))
            toc_cache(out_docx, pmap)
            if pdf_out:
                shutil.copy2(out_docx, tmp_src)
                tmp_final = os.path.join(tmp_dir, "final.pdf")
                if rdx.render_to_pdf(tmp_src, tmp_final):
                    shutil.copy2(tmp_final, pdf_out)
                    final_pdf = pdf_out
        shutil.rmtree(tmp_dir, ignore_errors=True)
    except Exception as exc:
        print("TOC cache skipped: %s" % exc, file=sys.stderr)
    sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "core"))
    from result_document_quality_engine import run as qrun
    report = qrun(out_docx, pdf_path=final_pdf, golden_path=golden)
    print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
