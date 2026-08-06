# -*- coding: utf-8 -*-
import io
import os

from docx import Document


def main():
    base = r"D:\Users\leibaihong\Desktop\课程材料优化\毕业设计智能制作工作区\06_输出成果\03_电梯系统\答辩差异化批量验证"
    p = os.path.join(base, "邱志豪_成果包",
                     "04 邱志豪 毕业设计答辩记录表 太平街口站电梯常见故障分析与检修方案设计.docx")
    out = []
    doc = Document(p)
    cell = doc.tables[0].rows[0].cells[0]
    out.append("== generated cell paragraphs ==")
    for i, pp in enumerate(cell.paragraphs):
        out.append("%02d %r" % (i, pp.text[:50]))

    src = r"D:\Users\leibaihong\Desktop\课程材料优化\毕业设计智能制作工作区\03_需要修改文件整理\03_电梯系统\邱志豪\答辩记录表.docx"
    sdoc = Document(src)
    if len(sdoc.tables) >= 2:
        scell = sdoc.tables[1].rows[0].cells[0]
        out.append("== source table[1] paragraphs ==")
        for i, pp in enumerate(scell.paragraphs[:8]):
            out.append("%02d %r" % (i, pp.text[:60]))
    with io.open(os.path.join(os.path.dirname(os.path.abspath(__file__)), "_batch_cell_inspect.txt"),
                 "w", encoding="utf-8") as f:
        f.write("\n".join(out))
    print("saved")


if __name__ == "__main__":
    main()
