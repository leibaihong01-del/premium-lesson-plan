# -*- coding: utf-8 -*-
import io
import os
import sys

from docx import Document


def main():
    base = r"D:\Users\leibaihong\Desktop\课程材料优化\毕业设计智能制作工作区\06_输出成果\03_电梯系统\答辩差异化批量验证"
    students = {
        "邱志豪": ("202421044713", "24级机电技术2班", "太平街口站电梯常见故障分析与检修方案设计"),
        "李增伟": ("202421044707", "24级机电技术2班", "南郊公园站自动扶梯日常维护方案设计"),
        "谭子安": ("202421044727", "24级机电技术2班", "梧桐路站电梯门系统检修方案设计"),
    }
    out = []
    for name, (sid, cls, topic) in students.items():
        pkg = os.path.join(base, name + "_成果包")
        docx = os.path.join(pkg, "04 %s 毕业设计答辩记录表 %s.docx" % (name, topic))
        pdf = os.path.join(pkg, "_过程记录", "04 %s 毕业设计答辩记录表 %s.pdf" % (name, topic))
        if not os.path.exists(docx):
            out.append("%s: DOCX MISSING" % name)
            continue
        doc = Document(docx)
        texts = [pp.text for pp in doc.paragraphs]
        for t in doc.tables:
            for row in t.rows:
                for c in row.cells:
                    texts.append(c.text)
        full = "\n".join(texts)
        hits = {
            "name": name in full,
            "sid": sid in full,
            "class": cls in full,
            "topic": topic in full,
            "pdf": os.path.exists(pdf),
        }
        cell = doc.tables[0].rows[0].cells[0]
        paras = [pp.text.strip() for pp in cell.paragraphs if pp.text.strip()]
        hits["answer_lines"] = len(paras)
        hits["first_line"] = paras[0][:30] if paras else ""
        out.append("%s %s" % (name, hits))
    report = os.path.join(base, "defense_batch_differentiation_report.json")
    out.append("report exists: %s" % os.path.exists(report))
    with io.open(os.path.join(base, "batch_validation.txt"), "w", encoding="utf-8") as f:
        f.write("\n".join(out))
    print("saved validation")


if __name__ == "__main__":
    main()
