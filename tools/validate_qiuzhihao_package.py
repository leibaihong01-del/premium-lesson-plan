# -*- coding: utf-8 -*-
import io
import json
import os
import sys

import pdfplumber
from docx import Document

PACKAGE = sys.argv[1]
PROCESS = os.path.join(PACKAGE, "_过程记录")
NAME = "邱志豪"
SID = "202421044713"
CLS = "24级机电技术2班"
ADV = "瞿曌"
TOPIC = "太平街口站电梯常见故障分析与检修方案设计"
FILES = {
    "01": "01 邱志豪 毕业设计任务书 太平街口站电梯常见故障分析与检修方案设计.docx",
    "02": "02 邱志豪 毕业设计成果 太平街口站电梯常见故障分析与检修方案设计.docx",
    "03": "03 邱志豪 毕业设计成绩评定表 太平街口站电梯常见故障分析与检修方案设计.docx",
    "04": "04 邱志豪 毕业设计答辩记录表 太平街口站电梯常见故障分析与检修方案设计.docx",
}


def main():
    out = []
    for code, fn in FILES.items():
        p = os.path.join(PACKAGE, fn)
        if not os.path.exists(p):
            out.append("%s: MISSING" % code)
            continue
        doc = Document(p)
        texts = [para.text for para in doc.paragraphs]
        for t in doc.tables:
            for row in t.rows:
                for c in row.cells:
                    texts.append(c.text)
        full = "\n".join(texts)
        hits = {k: (v in full) for k, v in [
            ("name", NAME), ("sid", SID), ("class", CLS), ("advisor", ADV), ("topic", TOPIC)]}
        out.append("%s %s" % (code, hits))

    report = os.path.join(PROCESS, "result_production_report.json")
    if os.path.exists(report):
        with io.open(report, encoding="utf-8") as f:
            rp = json.load(f)
        out.append("result status: %s" % rp.get("status"))
        wq = rp.get("dimensions", {}).get("word_structure", {})
        out.append("result sections: %s footer_parity: %s toc_cache: %s" % (
            wq.get("sections"), wq.get("footer_page_parity"), wq.get("toc_cache_entries")))

    tb_pdf = os.path.join(PROCESS, "01 邱志豪 毕业设计任务书 太平街口站电梯常见故障分析与检修方案设计.pdf")
    if os.path.exists(tb_pdf):
        with pdfplumber.open(tb_pdf) as pdf:
            out.append("taskbook pdf pages: %d" % len(pdf.pages))

    out.append("== files ==")
    for root_dir, dirs, fnames in os.walk(PACKAGE):
        for fn in sorted(fnames):
            fp = os.path.join(root_dir, fn)
            out.append("%s (%d)" % (fp, os.path.getsize(fp)))

    with io.open(os.path.join(PROCESS, "package_validation.txt"), "w", encoding="utf-8") as f:
        f.write("\n".join(out))
    print("saved validation")


if __name__ == "__main__":
    main()
