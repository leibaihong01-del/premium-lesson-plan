# -*- coding: utf-8 -*-
import importlib.util
import os
import shutil
import sys
import tempfile

from docx import Document
from docx.oxml.ns import qn


def load_render(mod_path):
    spec = importlib.util.spec_from_file_location("render_docx", mod_path)
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod


def main():
    docx = sys.argv[1]
    pdf = sys.argv[2]
    render_mod = sys.argv[3]

    doc = Document(docx)
    t = doc.tables[0]
    for ri in (4, 5):
        cell = t.rows[ri].cells[1]
        empties = [p for p in cell.paragraphs if not p.text.strip()]
        while len(empties) > 2:
            p = empties.pop()
            p._p.getparent().remove(p._p)
        trPr = t.rows[ri]._tr.find(qn("w:trPr"))
        if trPr is None:
            trPr = t.rows[ri]._tr.makeelement(qn("w:trPr"), {})
            t.rows[ri]._tr.insert(0, trPr)
        h = trPr.find(qn("w:trHeight"))
        if h is None:
            h = trPr.makeelement(qn("w:trHeight"), {})
            trPr.append(h)
        h.set(qn("w:val"), "0")
        h.set(qn("w:hRule"), "auto")
    doc.save(docx)

    rd = load_render(render_mod)
    tmp = tempfile.mkdtemp(prefix="taskbook_fit_")
    src = os.path.join(tmp, "src.docx")
    dst = os.path.join(tmp, "out.pdf")
    shutil.copy2(docx, src)
    ok = rd.render_to_pdf(src, dst)
    pages = None
    if ok and os.path.exists(dst):
        try:
            from pypdf import PdfReader
            pages = len(PdfReader(dst).pages)
        except Exception:
            try:
                import pdfplumber
                with pdfplumber.open(dst) as p:
                    pages = len(p.pages)
            except Exception:
                pass
        os.makedirs(os.path.dirname(pdf), exist_ok=True)
        shutil.copy2(dst, pdf)
    shutil.rmtree(tmp, ignore_errors=True)
    print("render ok:", ok, "pages:", pages)


if __name__ == "__main__":
    main()
