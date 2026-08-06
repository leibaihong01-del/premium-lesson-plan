# -*- coding: utf-8 -*-
"""GitHub 历史 Result Skill 独立回归测试（戴吉祥）。"""
import json, os, re, subprocess, sys
sys.stdout.reconfigure(encoding="utf-8", errors="replace")
COURSEAGENT = r"D:\Users\leibaihong\Desktop\课程材料优化\CourseAgent"
PROJECT = r"D:\Users\leibaihong\Desktop\课程材料优化"
WS = os.path.join(PROJECT, "毕业设计智能制作工作区")
OUT_DOCS = os.path.join(COURSEAGENT, "docs", "v0.7", "GitHubHistorical")
os.makedirs(OUT_DOCS, exist_ok=True)
sys.path.insert(0, os.path.join(COURSEAGENT, "core"))
sys.path.insert(0, os.path.join(WS, "00_系统配置", "模块", "v03"))
sys.path.insert(0, os.path.join(WS, "00_系统配置", "模块", "v06"))
from docx import Document
import result_reference_builder as rrb
from graduation_skill_runners import _apply_reference_format_experience, _normalize_result_fonts, _transplant_cover
from render_docx import render_to_pdf
from result_document_parser import parse as parse_result

def git(*args):
    r = subprocess.run(["git", "-C", COURSEAGENT] + list(args), capture_output=True, text=True, encoding="utf-8", errors="replace")
    return r.stdout

# 第一步：GitHub 历史 Skill 索引
files = git("ls-tree", "-r", "--name-only", "v1.0-result-baseline").splitlines()
relevant = [f for f in files if any(k in f for k in ["skills/graduation_design", "prompts/", "modules/template_parser", "experience", "strategy", "memory", "template_schema"])]
index = []
for f in relevant:
    cat = "unknown"
    if f.startswith("prompts/"): cat = "prompt"
    elif "result/SKILL.md" in f: cat = "skill"
    elif "strategy" in f or "generation_rules" in f: cat = "strategy"
    elif "memory" in f or "experience" in f: cat = "experience"
    elif "template_schema" in f or "template_parser" in f: cat = "template"
    index.append({"category": cat, "path": f, "source": "tag:v1.0-result-baseline", "version": "v1.0-result-baseline"})
with open(os.path.join(OUT_DOCS, "Github_Result_Skill_Index.json"), "w", encoding="utf-8") as f:
    json.dump({"total": len(index), "assets": index}, f, ensure_ascii=False, indent=2)

# 第三步：生成戴吉祥 GitHub Skill 验证版
DIRECTION = "06_空调通风系统"
STUDENT = "戴吉祥"
INFO = os.path.join(WS, "03_需要修改文件整理", DIRECTION, STUDENT, "学生信息.json")
STUDENT_DIR = os.path.join(WS, "03_需要修改文件整理", DIRECTION, STUDENT)
OUT_DIR = os.path.join(WS, "06_输出成果", DIRECTION, STUDENT + "_GitHubSkill验证")
os.makedirs(OUT_DIR, exist_ok=True)
with open(INFO, encoding="utf-8") as f:
    info = json.load(f)
topic = info["课题名称"]
template = os.path.join(WS, "02_模板文件", "02 杨振海 毕业设计成果 黄兴南路站AFC闸机设备检修方案设计.docx")
output = os.path.join(OUT_DIR, "戴吉祥毕业设计成果_GitHubSkill验证版.docx")
pdf = os.path.join(OUT_DIR, "戴吉祥毕业设计成果_GitHubSkill验证版.pdf")
rrb.reconstruct(STUDENT, DIRECTION, "02", info, os.path.join(STUDENT_DIR, "成果初稿.docx"), output, neutralize=True)
doc = Document(output)
cover = os.path.join(STUDENT_DIR, "成果记录表.docx")
if os.path.isfile(cover):
    _transplant_cover(doc, Document(cover))
_normalize_result_fonts(doc)
_apply_reference_format_experience(doc)
for p in doc.paragraphs:
    t = p.text.strip()
    m = re.match(r"^表\s*(\d+)[-－](\d+)", t)
    if m:
        new = "表%s.%s  %s" % (m.group(1), m.group(2), t[m.end():].strip())
        if p.runs:
            p.runs[0].text = new
            for r in p.runs[1:]:
                r.text = ""
doc.save(output)
render_to_pdf(output, pdf)

def metrics(path):
    m = parse_result(path)
    return {"chars": m.get("body_chars"), "sections": len(m.get("sections", [])), "tables": len(m.get("tables", [])), "refs": m.get("reference_count", 0), "fonts": m.get("body_font_sizes"), "toc": m.get("toc_field_present")}
new = metrics(output)
wang = metrics(os.path.join(WS, "06_输出成果", "V0.4.2_王欢成果验证", "02 王欢 毕业设计成果 解放西路站AFC检票机故障排查方案设计_最终验收版.docx"))
qiu = metrics(os.path.join(WS, "06_输出成果", "V0.4.2_邱志豪成果验证", "01 邱志豪 毕业设计成果 太平街口站电梯常见故障分析与检修方案设计_最终验收版.docx"))

report = [
    "# GitHub Skill Regression Report", "",
    "- 学生：戴吉祥", "- Skill版本：v1.0-result-baseline", "- Commit：ed58bee / e30468f", "",
    "| 指标 | GitHub Skill版 | 王欢最终验收版 | 邱志豪最终验收版 |", "|---|---|---|---|",
    "| 正文字数 | %s | %s | %s |" % (new["chars"], wang["chars"], qiu["chars"]),
    "| 章节数 | %s | %s | %s |" % (new["sections"], wang["sections"], qiu["sections"]),
    "| 表格数 | %s | %s | %s |" % (new["tables"], wang["tables"], qiu["tables"]),
    "| 参考文献 | %s | %s | %s |" % (new["refs"], wang["refs"], qiu["refs"]),
    "| 正文字号 | %s | %s | %s |" % (new["fonts"], wang["fonts"], qiu["fonts"]),
    "| TOC | %s | %s | %s |" % (new["toc"], wang["toc"], qiu["toc"]),
    "", "## 生成链", "",
    "- 入口：result_reference_builder（GitHub Skill v1.0 声明的模板保真重构）",
    "- Prompt：prompts/manifest.yaml + prompts/translator/system.md",
    "- 依赖：template_schema.json / result_generation_strategy.yaml / audit_rules.json",
]
report_path = os.path.join(OUT_DOCS, "GitHub_Skill_Regression_Report.md")
with open(report_path, "w", encoding="utf-8") as f:
    f.write("\n".join(report))

print(json.dumps({"index": os.path.join(OUT_DOCS, "Github_Result_Skill_Index.json"), "docx": output, "pdf": pdf, "report": report_path, "metrics": {"github": new, "wang": wang, "qiu": qiu}}, ensure_ascii=False, indent=2))