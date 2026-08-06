# -*- coding: utf-8 -*-
import io
import os

from docx import Document


def main():
    base = r"D:\Users\leibaihong\Desktop\课程材料优化\毕业设计智能制作工作区\06_输出成果\03_电梯系统\答辩差异化批量验证"
    students = ["邱志豪", "李增伟", "谭子安"]
    topics = {
        "邱志豪": "太平街口站电梯常见故障分析与检修方案设计",
        "李增伟": "南郊公园站自动扶梯日常维护方案设计",
        "谭子安": "梧桐路站电梯门系统检修方案设计",
    }
    out = []
    for s in students:
        p = os.path.join(base, s + "_成果包",
                         "04 %s 毕业设计答辩记录表 %s.docx" % (s, topics[s]))
        doc = Document(p)
        cell = doc.tables[0].rows[0].cells[0]
        paras = [pp.text.strip() for pp in cell.paragraphs if pp.text.strip()]
        out.append("== %s ==" % s)
        for t in paras[:3]:
            out.append(t[:70])
    with io.open(os.path.join(os.path.dirname(os.path.abspath(__file__)), "_batch_first_sentences.txt"),
                 "w", encoding="utf-8") as f:
        f.write("\n".join(out))
    print("saved")


if __name__ == "__main__":
    main()
