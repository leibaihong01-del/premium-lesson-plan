# -*- coding: utf-8 -*-
"""Document DNA：优秀成果对象级逆向解析（只读）。"""
import json, os, re, sys
sys.stdout.reconfigure(encoding="utf-8", errors="replace")
from docx import Document
from docx.oxml.ns import qn

DOCX = r"D:\Users\leibaihong\Desktop\课程材料优化\毕业设计智能制作工作区\06_输出成果\V0.4.2_邱志豪成果验证\01 邱志豪 毕业设计成果 太平街口站电梯常见故障分析与检修方案设计_最终验收版.docx"
PDF = r"D:\Users\leibaihong\Desktop\课程材料优化\毕业设计智能制作工作区\06_输出成果\V0.4.2_渲染验收\邱志豪_最终版.pdf"
OUT = r"D:\Users\leibaihong\Desktop\课程材料优化\CourseAgent\docs\v0.7\DocumentDNA"
os.makedirs(OUT, exist_ok=True)

def emu_cm(v):
    return round(v / 360000.0, 2) if v is not None else None

def tw_cm(v):
    return round(v / 567.0, 2) if v is not None else None

def run_info(r):
    rPr = r._r.find(qn("w:rPr"))
    east = None
    if rPr is not None:
        rf = rPr.find(qn("w:rFonts"))
        if rf is not None:
            east = rf.get(qn("w:eastAsia"))
    color = None
    if rPr is not None:
        c = rPr.find(qn("w:color"))
        if c is not None:
            color = c.get(qn("w:val"))
    return {
        "text": r.text,
        "font": r.font.name,
        "east_asia": east,
        "size_pt": r.font.size.pt if r.font.size else None,
        "bold": bool(r.font.bold),
        "italic": bool(r.font.italic),
        "underline": bool(r.font.underline),
        "color": color,
    }

def para_info(p, idx):
    pf = p.paragraph_format
    return {
        "index": idx,
        "style": p.style.name if p.style else "",
        "text": p.text,
        "chars": len(p.text),
        "alignment": str(p.alignment) if p.alignment is not None else None,
        "line_spacing": pf.line_spacing,
        "space_before": pf.space_before,
        "space_after": pf.space_after,
        "left_indent": pf.left_indent,
        "first_line_indent": pf.first_line_indent,
        "page_break_before": bool(pf.page_break_before),
        "keep_with_next": bool(pf.keep_with_next),
        "runs": [run_info(r) for r in p.runs if r.text],
    }

def role_of(p):
    t = p.text.strip()
    style = p.style.name if p.style else ""
    if "Heading 1" in style or re.match(r"^\d+\s+\S", t):
        return "heading1"
    if "Heading 2" in style or re.match(r"^\d+\.\d+\s+\S", t):
        return "heading2"
    if "Heading 3" in style or re.match(r"^\d+\.\d+\.\d+\s+\S", t) or re.match(r"^（[一二三四五六七八九十]+）", t):
        return "heading3"
    if "参考文献" in style or t == "参考文献":
        return "reference_heading" if t == "参考文献" else "reference"
    if "表注" in style or re.match(r"^表\s*\d", t):
        return "caption"
    return "body"

doc = Document(DOCX)
paras = []
for i, p in enumerate(doc.paragraphs):
    pi = para_info(p, i)
    pi["role"] = role_of(p)
    paras.append(pi)

# sections
sections = []
for i, s in enumerate(doc.sections):
    sections.append({
        "index": i,
        "page_width_cm": emu_cm(s.page_width),
        "page_height_cm": emu_cm(s.page_height),
        "margins_cm": {"top": emu_cm(s.top_margin), "right": emu_cm(s.right_margin), "bottom": emu_cm(s.bottom_margin), "left": emu_cm(s.left_margin)},
        "header": [p.text for p in s.header.paragraphs if p.text.strip()],
        "footer": [p.text for p in s.footer.paragraphs if p.text.strip()],
    })

# styles
style_dna = {}
for st in doc.styles:
    try:
        f = st.font
        pf = st.paragraph_format
        style_dna[st.name] = {
            "type": st.type,
            "base": st.base_style.name if st.base_style else None,
            "font_name": f.name,
            "size_pt": f.size.pt if f.size else None,
            "bold": bool(f.bold),
            "italic": bool(f.italic),
            "alignment": str(pf.alignment) if pf.alignment is not None else None,
            "line_spacing": pf.line_spacing,
            "space_before": pf.space_before,
            "space_after": pf.space_after,
            "left_indent": pf.left_indent,
            "first_line_indent": pf.first_line_indent,
        }
    except Exception:
        pass

# character scan
char_issues = {"chinese_font_anomalies": [], "halfwidth_punct_near_cjk": 0, "double_spaces": 0, "tabs": 0, "control_chars": 0}
CJK = re.compile(r"[\u4e00-\u9fff]")
for idx, p in enumerate(doc.paragraphs):
    text = p.text
    if re.search(r" {2,}", text):
        char_issues["double_spaces"] += 1
    if "\t" in text:
        char_issues["tabs"] += 1
    for ch in text:
        if ord(ch) < 32 and ch not in "\t\n\r":
            char_issues["control_chars"] += 1
    for r in p.runs:
        rPr = r._r.find(qn("w:rPr"))
        east = None
        if rPr is not None:
            rf = rPr.find(qn("w:rFonts"))
            if rf is not None:
                east = rf.get(qn("w:eastAsia"))
        if CJK.search(r.text) and east and east not in ("宋体",):
            char_issues["chinese_font_anomalies"].append({"paragraph": idx, "font": east, "sample": r.text[:20]})
    for ch in text:
        if ch in "\"',().;" and CJK.search(text[max(0,text.index(ch)-1):text.index(ch)+1]):
            char_issues["halfwidth_punct_near_cjk"] += 1

# references
refs = []
in_ref = False
for p in paras:
    if p["text"].strip() == "参考文献":
        in_ref = True
        continue
    if in_ref and p["text"].strip():
        refs.append({"paragraph": p["index"], "text": p["text"], "style": p["style"], "left_indent": p["left_indent"], "first_line_indent": p["first_line_indent"], "runs": p["runs"]})

# TOC
toc_instr = []
for instr in doc.element.body.iter(qn("w:instrText")):
    if instr.text and "TOC" in instr.text:
        toc_instr.append(instr.text)
toc_styles = {name: style_dna[name] for name in style_dna if re.match(r"TOC\s?\d", name)}

# figures
figures = []
for child in doc.element.body.iterchildren():
    if child.tag != qn("w:p"):
        continue
    drawings = list(child.iter(qn("w:drawing"))) + list(child.iter(qn("w:pict")))
    if drawings:
        ext = []
        for d in drawings:
            for e in d.iter(qn("wp:extent")):
                ext.append({"cx": e.get("cx"), "cy": e.get("cy")})
        figures.append({"type": "drawing" if drawings else "pict", "extents": ext[:2]})

# tables object-level
table_dna = []
body_children = list(doc.element.body.iterchildren())
last_para = None
tbl_idx = 0
for child in body_children:
    if child.tag == qn("w:p"):
        from docx.text.paragraph import Paragraph
        last_para = Paragraph(child, doc)
    elif child.tag == qn("w:tbl"):
        t = doc.tables[tbl_idx]
        caption = None
        caption_info = None
        if last_para is not None:
            txt = last_para.text.strip()
            if re.match(r"^表\s*\d", txt):
                caption = txt
                caption_info = para_info(last_para, -1)
        grid = [tw_cm(int(gc.get(qn("w:w"))) if gc.get(qn("w:w")) else 0) for gc in t._tbl.findall(qn("w:tblGrid"))]
        rows = []
        for ri, row in enumerate(t.rows):
            trPr = row._tr.find(qn("w:trPr"))
            repeat_header = False
            cant_split = False
            row_height = None
            if trPr is not None:
                repeat_header = trPr.find(qn("w:tblHeader")) is not None
                cant_split = trPr.find(qn("w:cantSplit")) is not None
                th = trPr.find(qn("w:trHeight"))
                if th is not None:
                    row_height = th.get(qn("w:val"))
            cells = []
            seen = set()
            for ci, c in enumerate(row.cells):
                if c._tc in seen:
                    continue
                seen.add(c._tc)
                tcPr = c._tc.find(qn("w:tcPr"))
                v_align = None
                tc_w = None
                if tcPr is not None:
                    va = tcPr.find(qn("w:vAlign"))
                    if va is not None:
                        v_align = va.get(qn("w:val"))
                    tcW = tcPr.find(qn("w:tcW"))
                    if tcW is not None:
                        tc_w = tcW.get(qn("w:w"))
                cell_paras = [{"text": pp.text, "style": pp.style.name if pp.style else "", "runs": [run_info(r) for r in pp.runs if r.text]} for pp in c.paragraphs]
                cells.append({"col": ci, "text": c.text, "width_twips": tc_w, "v_align": v_align, "paragraphs": cell_paras})
            rows.append({"row": ri, "header_row": ri == 0, "repeat_header": repeat_header, "cant_split": cant_split, "height_twips": row_height, "cells": cells})
        table_dna.append({
            "table_id": tbl_idx,
            "rows": len(t.rows),
            "cols": len(t.columns),
            "caption": caption,
            "caption_paragraph": caption_info,
            "grid_width_cm": grid,
            "rows_data": rows,
        })
        tbl_idx += 1

# content architecture
chapters = []
current = None
for p in paras:
    if p["role"] == "heading1":
        current = {"title": p["text"], "body_chars": 0, "tables": 0, "captions": 0}
        chapters.append(current)
    elif current is not None and p["role"] == "body":
        current["body_chars"] += p["chars"]
    elif current is not None and p["role"] == "caption":
        current["captions"] += 1

# PDF visual
pdf_visual = {"page_count": None, "pages": []}
if os.path.isfile(PDF):
    import pdfplumber
    with pdfplumber.open(PDF) as pdf:
        pdf_visual["page_count"] = len(pdf.pages)
        for i, page in enumerate(pdf.pages):
            words = page.extract_words()
            text_bbox = None
            if words:
                text_bbox = {"x0": round(min(w["x0"] for w in words),1), "x1": round(max(w["x1"] for w in words),1), "top": round(min(w["top"] for w in words),1), "bottom": round(max(w["bottom"] for w in words),1)}
            pdf_visual["pages"].append({
                "page": i+1,
                "chars": len(page.chars),
                "tables": len(page.find_tables()),
                "images": len(page.images),
                "text_bbox": text_bbox,
            })

visual_baseline = {"source": PDF, "page_count": pdf_visual["page_count"], "pages": pdf_visual["pages"]}

# write outputs
def w(name, data):
    with open(os.path.join(OUT, name), "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

w("Document_Structure.json", {"sections": sections, "paragraph_count": len(paras), "paragraphs": paras})
w("Style_DNA.json", style_dna)
w("Character_Rule.json", {"rule": "中文宋体/英文Times New Roman；中文语境全角标点；禁止多余空格与隐藏字符", "issues": char_issues})
w("Reference_Rule.json", {"count": len(refs), "items": refs, "rule": "编号[1]开始连续；多行参考文献应保持悬挂缩进"})
w("TOC_Rule.json", {"toc_fields": toc_instr, "toc_styles": toc_styles, "rule": "正文标题变化后更新TOC域"})
w("Figure_Rule.json", {"figures": figures, "rule": "图注与图同页绑定"})
w("Table_DNA.json", table_dna)
w("Table_Caption_Rule.json", {"rule": "表注位于表格正上方，与表格同页；表注格式 表X.X + 两空格", "captions": [{"table": t["table_id"], "caption": t["caption"], "caption_paragraph": t["caption_paragraph"]} for t in table_dna]})
w("Table_Width_Model.json", {"tables": [{"table_id": t["table_id"], "columns_cm": t["grid_width_cm"]} for t in table_dna]})
w("Page_Visual_Fingerprint.json", pdf_visual)
w("Visual_Baseline.json", visual_baseline)
w("Content_Architecture.json", {"chapters": chapters, "references": len(refs), "tables": len(table_dna)})
w("Heading_Rules.json", {"headings": [{"index": p["index"], "text": p["text"], "role": p["role"], "style": p["style"]} for p in paras if p["role"].startswith("heading")]})

constraints = []
constraints.append("# Document Generation Constraint")
constraints.append("")
constraints.append("- 分节：%s；页宽/页高：%s/%s cm" % (len(sections), sections[0]["page_width_cm"] if sections else "?", sections[0]["page_height_cm"] if sections else "?"))
constraints.append("- 正文字体：宋体 12pt；英文/数字：Times New Roman")
constraints.append("- 标题：Heading 1/2/3，一级分页")
constraints.append("- 表注：表X.X + 两空格，位于表格正上方并同页")
constraints.append("- 表格：禁止重建，表头跨页重复")
constraints.append("- 参考文献：[N]连续编号，多行保持悬挂缩进")
constraints.append("- 目录：TOC域，正文标题变化后更新")
constraints.append("- 图片/图注：图注与图同页绑定")
constraints.append("- 内容深度：正文建议不少于5000字；章节覆盖背景/概述/故障/方案/总结")
with open(os.path.join(OUT, "Document_Generation_Constraint.md"), "w", encoding="utf-8") as f:
    f.write("\n".join(constraints))

print(json.dumps({
    "out_dir": OUT,
    "paragraphs": len(paras),
    "sections": len(sections),
    "tables": len(table_dna),
    "references": len(refs),
    "figures": len(figures),
    "chapters": len(chapters),
    "pdf_pages": pdf_visual["page_count"],
}, ensure_ascii=False, indent=2))