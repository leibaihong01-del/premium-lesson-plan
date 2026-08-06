# -*- coding: utf-8 -*-
"""GitHub 历史成果 Skill 恢复与王欢复测。"""
import json, os, subprocess, sys
sys.stdout.reconfigure(encoding="utf-8", errors="replace")
COURSEAGENT = r"D:\Users\leibaihong\Desktop\课程材料优化\CourseAgent"
PROJECT = r"D:\Users\leibaihong\Desktop\课程材料优化"
WS = os.path.join(PROJECT, "毕业设计智能制作工作区")
sys.path.insert(0, os.path.join(COURSEAGENT, "core"))
sys.path.insert(0, os.path.join(WS, "00_系统配置", "模块", "v03"))
sys.path.insert(0, os.path.join(WS, "00_系统配置", "模块", "v06"))
from docx import Document
import result_reference_builder as rrb
from graduation_skill_runners import _apply_reference_format_experience, _normalize_result_fonts, _transplant_cover
from render_docx import render_to_pdf
from result_document_parser import parse as parse_result

OUT_DOCS = os.path.join(COURSEAGENT, "docs", "v0.7", "GitHubHistorical")
os.makedirs(OUT_DOCS, exist_ok=True)
repo = COURSEAGENT
def git(*args):
    r = subprocess.run(["git", "-C", repo] + list(args), capture_output=True, text=True, encoding="utf-8", errors="replace")
    return r.stdout

files = git("ls-tree", "-r", "--name-only", "v1.0-result-baseline").splitlines()
relevant = [f for f in files if any(k in f for k in ["skills/graduation_design/result", "prompts/", "modules/template_parser", "experience", "strategy", "template_schema", "result_generation_strategy", "quality_rule"])]
index = []
for f in relevant:
    cat = "unknown"
    if f.startswith("prompts/"): cat = "prompt"
    elif "result/SKILL.md" in f: cat = "skill"
    elif "strategy" in f or "generation_rules" in f: cat = "strategy"
    elif "experience" in f or "memory" in f or "candidate" in f: cat = "experience"
    elif "template_schema" in f or "template_mapping" in f: cat = "template"
    elif "rule" in f or "quality" in f or "audit" in f: cat = "quality"
    index.append({"category": cat, "path": f, "source": "tag:v1.0-result-baseline", "version": "v1.0-result-baseline"})
with open(os.path.join(OUT_DOCS, "GitHub_Historical_Skill_Index.json"), "w", encoding="utf-8") as f:
    json.dump({"total": len(index), "assets": index}, f, ensure_ascii=False, indent=2)

chain = {
    "schema": "0.7-gh-historical-chain-v1",
    "skill_version": "v1.0-result-baseline",
    "commit": ["ed58bee", "e30468f"],
    "prompt_version": "prompts/manifest.yaml + prompts/translator/system.md（tag v1.0-result-baseline）",
    "template_version": "02 杨振海 毕业设计成果 黄兴南路站AFC闸机设备检修方案设计.docx",
    "experience_sources": ["skills/graduation_design/result/memory/golden_cases/wanghuan.md", "skills/graduation_design/result/experience/candidates"],
    "generation_entry": "result_reference_builder（模板保真重构）",
    "quality_gates": ["result/SKILL.md", "quality_rule.yaml", "audit_rules.json"]
}
with open(os.path.join(OUT_DOCS, "Historical_Result_Generation_Chain.json"), "w", encoding="utf-8") as f:
    json.dump(chain, f, ensure_ascii=False, indent=2)

# 王欢 GitHub 历史链复测
DIRECTION = "01_AFC自动售检票系统"
STUDENT = "王欢"
INFO = os.path.join(WS, "03_需要修改文件整理", DIRECTION, STUDENT, "学生信息.json")
STUDENT_DIR = os.path.join(WS, "03_需要修改文件整理", DIRECTION, STUDENT)
OUT_DIR = os.path.join(WS, "06_输出成果", DIRECTION, STUDENT + "_GitHub历史链复测")
os.makedirs(OUT_DIR, exist_ok=True)
with open(INFO, encoding="utf-8") as f:
    info = json.load(f)
topic = info["课题名称"]
template = os.path.join(WS, "02_模板文件", "02 杨振海 毕业设计成果 黄兴南路站AFC闸机设备检修方案设计.docx")
output = os.path.join(OUT_DIR, "王欢毕业设计成果_GitHub历史链复测版.docx")
pdf = os.path.join(OUT_DIR, "王欢毕业设计成果_GitHub历史链复测版.pdf")
rrb.reconstruct(STUDENT, DIRECTION, "02", info, os.path.join(STUDENT_DIR, "成果初稿.docx"), output, neutralize=True)
doc = Document(output)
cover = os.path.join(STUDENT_DIR, "成果记录表.docx")
if os.path.isfile(cover):
    _transplant_cover(doc, Document(cover))
_normalize_result_fonts(doc)
_apply_reference_format_experience(doc)
for p in doc.paragraphs:
    t = p.text.strip()
    m = __import__("re").match(r"^表\s*(\d+)[-－](\d+)", t)
    if m:
        new = "表%s.%s  %s" % (m.group(1), m.group(2), t[m.end():].strip())
        if p.runs:
            p.runs[0].text = new
            for r in p.runs[1:]: r.text = ""
doc.save(output)
render_to_pdf(output, pdf)

new_m = parse_result(output)
old_path = os.path.join(WS, "06_输出成果", "V0.4.2_王欢成果验证", "02 王欢 毕业设计成果 解放西路站AFC检票机故障排查方案设计_最终验收版.docx")
old_m = parse_result(old_path)
def metrics(m):
    return {"chars": m.get("body_chars"), "sections": len(m.get("sections", [])), "tables": len(m.get("tables", [])), "refs": m.get("reference_count", 0), "toc": m.get("toc_field_present")}
print(json.dumps({"index": os.path.join(OUT_DOCS, "GitHub_Historical_Skill_Index.json"), "chain": os.path.join(OUT_DOCS, "Historical_Result_Generation_Chain.json"), "docx": output, "pdf": pdf, "old_final": metrics(old_m), "new_github": metrics(new_m)}, ensure_ascii=False, indent=2))