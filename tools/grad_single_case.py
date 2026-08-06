# -*- coding: utf-8 -*-
"""GraduationAgent 单案例生成（正式模板 + 扁平统一命名 + 真实段落）。"""
import copy, json, os, re, time
from docx import Document

WS = os.environ.get("GRAD_WS", r"D:\Users\leibaihong\Desktop\课程材料优化\毕业设计智能管理工作区")
STUDENT_NO = "202421044606"
NAME = "唐维泽"
TOPIC = "香樟路站AFC闸机防夹功能检测方案设计"
TEACHER = "瞿曌"
CLASS_ = "24级机电技术1班"
MAJOR = "城市轨道车辆应用技术"
DIRECTION = "方向01"
PROJECT_DIR = NAME + "_" + STUDENT_NO

trace = []

def log(step, ok, msg=""):
    trace.append({"step": step, "ok": ok, "message": msg})
    print(("OK  " if ok else "FAIL"), step, msg)

def set_cell_text(cell, text):
    p = cell.paragraphs[0]
    runs = p.runs
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            r.text = ""
    else:
        p.add_run(text)

def set_cell_lines(cell, lines):
    ps = cell.paragraphs
    while len(ps) < len(lines):
        newp = copy.deepcopy(ps[-1]._p)
        ps[-1]._p.addnext(newp)
        ps = cell.paragraphs
    for extra in ps[len(lines):]:
        extra._p.getparent().remove(extra._p)
    for i, line in enumerate(lines):
        p = cell.paragraphs[i]
        runs = p.runs
        if runs:
            runs[0].text = line
            for r in runs[1:]:
                r.text = ""
        else:
            p.add_run(line)

def find_row_with_label(table, label):
    for i, row in enumerate(table.rows):
        seen = set()
        for c in row.cells:
            if c._tc in seen:
                continue
            seen.add(c._tc)
            if label in c.text:
                return i
    return -1

def set_cell_by_label(table, label, value, col=1):
    r = find_row_with_label(table, label)
    if r >= 0 and col < len(table.rows[r].cells):
        set_cell_text(table.rows[r].cells[col], value)
        return True
    return False

def fill_pairs(table, row, pairs):
    for col, val in pairs:
        set_cell_text(table.rows[row].cells[col], val)

def get_template(ws_root, tid):
    with open(os.path.join(ws_root, "00_系统配置", "template_index.json"), encoding="utf-8") as f:
        idx = json.load(f)
    for t in idx["templates"]:
        if t["id"] == tid:
            return os.path.join(ws_root, t["file"])
    return None

def main():
    base = os.path.join(WS, "outputs", DIRECTION, PROJECT_DIR)
    os.makedirs(base, exist_ok=True)
    log("init_project", True, base)

    # 1 taskbook
    try:
        src = get_template(WS, "taskbook")
        out = os.path.join(base, "01_毕业设计任务书.docx")
        doc = Document(src); t = doc.tables[0]
        fill_pairs(t, 0, [(1, "轨道车辆学院"), (3, NAME), (5, TEACHER)])
        fill_pairs(t, 1, [(1, CLASS_), (3, STUDENT_NO), (5, "")])
        fill_pairs(t, 2, [(1, TOPIC)])
        fill_pairs(t, 3, [(1, "方案设计类")])
        goal = ("深入理解香樟路站AFC闸机设备的结构、工作原理与检修要求；分析闸机防夹功能检测要点；"
                "制定科学合理的检测方案，提高设备运行可靠性与乘客通行安全。")
        set_cell_text(t.rows[4].cells[1], goal)
        set_cell_lines(t.rows[5].cells[1], [
            "1. 收集香樟路站AFC闸机设备资料与运行数据；",
            "2. 分析闸机防夹功能原理与检测要求；",
            "3. 设计防夹功能检测流程与判定标准；",
            "4. 形成完整检修方案并给出实施建议。",
        ])
        doc.save(out); log("taskbook_generate", True, out)
    except Exception as e:
        log("taskbook_generate", False, repr(e))

    # 2 result
    try:
        src = get_template(WS, "result")
        out = os.path.join(base, "02_毕业设计成果.docx")
        doc = Document(src); t = doc.tables[0]
        cover = {"选题名称": TOPIC, "选题类型": "方案设计类", "二级学院": "轨道车辆学院",
                 "专业名称": MAJOR, "班级名称": CLASS_, "姓名": NAME, "学号": STUDENT_NO,
                 "指导教师": TEACHER, "完成时间": "2026年10月"}
        for k, v in cover.items():
            set_cell_by_label(t, k, v, 1)
        draft = os.path.join(WS, "04_毕业设计成果", "原始草稿", "毕业设计成果初稿", "01 AFC.docx")
        body = []
        if os.path.exists(draft):
            dd = Document(draft)
            lines = [p.text.strip() for p in dd.paragraphs if p.text.strip()]
            idx = next((i for i, ln in enumerate(lines) if "防夹" in ln), -1)
            if idx >= 0:
                start = idx
                for j in range(idx, -1, -1):
                    if lines[j].startswith("1 引言"):
                        start = j; break
                end = len(lines)
                for j in range(idx + 1, len(lines)):
                    if lines[j].startswith("1 引言"):
                        end = j; break
                body = lines[start:end]
            else:
                body = lines[:80]
        bi = 0
        for p in doc.paragraphs:
            txt = p.text.strip()
            if bi >= len(body):
                break
            if re.match(r'^(一级标题|二级标题|三级标题|正文XX|参考文献|附\s*录|\d+\.\d+\s+XX)', txt) or txt == "一级标题1":
                if p.runs:
                    p.runs[0].text = body[bi]
                    for r in p.runs[1:]:
                        r.text = ""
                else:
                    p.add_run(body[bi])
                bi += 1
        doc.save(out)
        with open(os.path.join(base, "成果结构.json"), "w", encoding="utf-8") as f:
            json.dump({"topic": TOPIC, "body_lines": len(body)}, f, ensure_ascii=False, indent=2)
        log("result_generate", True, out + " body_lines=" + str(len(body)))
    except Exception as e:
        log("result_generate", False, repr(e))

    # 3 guidance
    try:
        src = get_template(WS, "guidance")
        out = os.path.join(base, "03_指导记录.docx")
        doc = Document(src); t = doc.tables[0]
        fill_pairs(t, 0, [(1, "轨道车辆学院"), (3, "57人"), (5, TEACHER)])
        fill_pairs(t, 1, [(1, "24级机电技术1-4班"), (3, "线上 线下"), (5, "")])
        doc.save(out); log("guidance_generate", True, out)
    except Exception as e:
        log("guidance_generate", False, repr(e))

    # 4 plagiarism placeholder
    md = os.path.join(base, "04_查重报告.md")
    js = os.path.join(base, "04_查重报告.json")
    with open(md, "w", encoding="utf-8") as f:
        f.write("# 查重报告（接口占位）\n\n- 学生：%s\n- 学号：%s\n- 查重率：待检测\n- 平台：待确认\n" % (NAME, STUDENT_NO))
    with open(js, "w", encoding="utf-8") as f:
        json.dump({"student_no": STUDENT_NO, "name": NAME, "duplication_rate": None,
                   "platform": None, "detected_at": None, "high_dup_sections": [],
                   "suggestions": [], "status": "pending"}, f, ensure_ascii=False, indent=2)
    log("plagiarism_placeholder", True, "接口预留")

    # 5 defense + score
    try:
        src = get_template(WS, "defense")
        out = os.path.join(base, "05_答辩及成绩评定.docx")
        doc = Document(src); t = doc.tables[0]
        fill_pairs(t, 0, [(1, "轨道车辆学院"), (3, NAME), (5, TEACHER)])
        fill_pairs(t, 1, [(1, CLASS_), (3, STUDENT_NO), (5, "")])
        fill_pairs(t, 2, [(1, TOPIC)])
        fill_pairs(t, 3, [(1, "方案设计类")])
        doc.save(out); log("defense_score_generate", True, out)
    except Exception as e:
        log("defense_score_generate", False, repr(e))

    trace_dir = os.path.join(WS, "00_系统配置", "运行轨迹")
    os.makedirs(trace_dir, exist_ok=True)
    with open(os.path.join(trace_dir, NAME + "_轨迹.json"), "w", encoding="utf-8") as f:
        json.dump(trace, f, ensure_ascii=False, indent=2)
    print("TRACE_DONE", sum(1 for x in trace if x["ok"]), "/", len(trace))

if __name__ == "__main__":
    main()