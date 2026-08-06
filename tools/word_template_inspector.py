# -*- coding: utf-8 -*-
"""Word 模板解析模块：提取文档结构、样式、表格与字段位置，生成模板规范 JSON。"""
import json
import os


def inspect_docx(path, max_paragraphs=150, max_rows=20):
    from docx import Document

    doc = Document(path)
    spec = {
        "file": os.path.basename(path),
        "sections": [],
        "paragraphs": [],
        "tables": [],
        "headers_footers": {"header": [], "footer": []},
    }
    for sec in doc.sections:
        spec["sections"].append({
            "page_width": sec.page_width,
            "page_height": sec.page_height,
            "orientation": str(sec.orientation),
            "left_margin": sec.left_margin,
            "right_margin": sec.right_margin,
            "top_margin": sec.top_margin,
            "bottom_margin": sec.bottom_margin,
            "header_distance": sec.header_distance,
            "footer_distance": sec.footer_distance,
            "start_type": str(sec.start_type),
        })
    for p in doc.paragraphs[:max_paragraphs]:
        runs = p.runs[:3]
        spec["paragraphs"].append({
            "style": p.style.name if p.style else None,
            "text": p.text.strip()[:60],
            "alignment": str(p.alignment),
            "first_line_indent": str(p.paragraph_format.first_line_indent),
            "line_spacing": str(p.paragraph_format.line_spacing),
            "runs": [{
                "text": r.text[:20],
                "bold": r.bold,
                "size_pt": r.font.size.pt if r.font.size else None,
                "font": r.font.name,
            } for r in runs],
        })
    for ti, t in enumerate(doc.tables):
        tbl = {
            "index": ti,
            "rows": len(t.rows),
            "cols": len(t.columns),
            "merge_signature": [],
            "rows_sample": [],
        }
        seen = set()
        for row in t.rows[: max_rows if max_rows else len(t.rows)]:
            sig = []
            cells = []
            for c in row.cells:
                if c._tc in seen:
                    sig.append("M")
                else:
                    seen.add(c._tc)
                    sig.append("U")
                    cells.append(c.text.strip().replace("\n", " / ")[:50])
            tbl["merge_signature"].append("".join(sig))
            tbl["rows_sample"].append(cells)
        spec["tables"].append(tbl)
    if doc.sections:
        for hf in ("header", "footer"):
            part = getattr(doc.sections[0], hf)
            spec["headers_footers"][hf] = [p.text.strip() for p in part.paragraphs if p.text.strip()]
    return spec


if __name__ == "__main__":
    import argparse
    ap = argparse.ArgumentParser()
    ap.add_argument("path")
    ap.add_argument("--out", default="")
    args = ap.parse_args()
    spec = inspect_docx(args.path)
    if args.out:
        with open(args.out, "w", encoding="utf-8") as f:
            json.dump(spec, f, ensure_ascii=False, indent=2)
        print("written", args.out)
    else:
        print(json.dumps(spec, ensure_ascii=False, indent=2))