# -*- coding: utf-8 -*-
"""电梯方向答辩差异化验证：三级画像 + 相似度 + 安全重写。"""
import io
import json
import os
import re
import sys

from docx import Document


def answer_text(path):
    doc = Document(path)
    if len(doc.tables) < 2:
        return ""
    cell = doc.tables[1].rows[0].cells[0]
    return "\n".join(p.text for p in cell.paragraphs)


def station_of(topic):
    m = re.search(r"([\u4e00-\u9fa5]{2,8}?站)", topic or "")
    return m.group(1) if m else ""


def main():
    base = r"D:\Users\leibaihong\Desktop\课程材料优化\毕业设计智能制作工作区\03_需要修改文件整理\03_电梯系统"
    students = ["邱志豪", "李增伟", "谭子安"]
    topics = {
        "邱志豪": "太平街口站电梯常见故障分析与检修方案设计",
        "李增伟": "南郊公园站自动扶梯日常维护方案设计",
        "谭子安": "梧桐路站电梯门系统检修方案设计",
    }
    sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "core"))
    import defense_differentiation as dd

    items = []
    for s in students:
        p = os.path.join(base, s, "答辩记录表.docx")
        ans = answer_text(p)
        topic = topics.get(s, "")
        items.append({
            "student": s,
            "professional": "城市轨道交通机电技术",
            "direction": "电梯系统",
            "topic": topic,
            "station": station_of(topic),
            "answer": ans,
        })

    profiles = dd.load_profiles()
    report = dd.process_batch(items, profiles=profiles)
    out = {
        "direction": "电梯系统",
        "students": students,
        "metrics": report["metrics"],
        "rewrites": report["rewrites"],
        "kept": report["kept"],
        "profiles": {
            "professional": dd.get_professional_profile("城市轨道交通机电技术", profiles),
            "direction": dd.build_direction_profile("城市轨道交通机电技术", "电梯系统", profiles=profiles),
        },
    }
    out_path = os.path.join(
        r"D:\Users\leibaihong\Desktop\课程材料优化\CourseAgent\docs\graduation_design",
        "defense_differentiation_elevator_validation.json")
    with io.open(out_path, "w", encoding="utf-8") as f:
        json.dump(out, f, ensure_ascii=False, indent=2)
    print("direction:", "电梯系统")
    print("metrics:", report["metrics"].get("电梯系统"))
    print("kept:", report["kept"])
    print("rewrites:", list(report["rewrites"].keys()))


if __name__ == "__main__":
    main()
