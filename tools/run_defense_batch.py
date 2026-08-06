# -*- coding: utf-8 -*-
"""同方向答辩相似度分析（只读，不生成、不修改 DOCX）。

答辩记录生成固定使用 DefenseSkillRunner + DEFENSE_LAYOUT_NORMALIZE=1（v0.9）。
本工具仅做相似度检测，供人工决定是否需要差异化，不写任何成果文件。
"""
import io
import json
import os
import re
import sys

from docx import Document


ROOT = r"D:\Users\leibaihong\Desktop\课程材料优化"
WS = os.path.join(ROOT, "毕业设计智能制作工作区")
CA = os.path.join(ROOT, "CourseAgent")
DIRECTION_DIR = "03_电梯系统"
DIRECTION = "电梯系统"
STUDENTS = [
    ("邱志豪", "太平街口站电梯常见故障分析与检修方案设计"),
    ("李增伟", "南郊公园站自动扶梯日常维护方案设计"),
    ("谭子安", "梧桐路站电梯门系统检修方案设计"),
]
OUT_DIR = os.path.join(WS, "05_质量检查", "答辩差异化分析")


def source_answer(student):
    p = os.path.join(WS, "03_需要修改文件整理", DIRECTION_DIR, student, "答辩记录表.docx")
    doc = Document(p)
    if len(doc.tables) < 2:
        return ""
    cell = doc.tables[1].rows[0].cells[0]
    return "\n".join(pp.text for pp in cell.paragraphs)


def main():
    sys.path.insert(0, os.path.join(CA, "core"))
    import defense_differentiation as dd

    texts = [source_answer(s) for s, _ in STUDENTS]
    profiles = dd.load_profiles()
    metrics = dd.analyze_similarity(texts, direction=DIRECTION, profiles=profiles)
    report = {
        "mode": "analysis-only",
        "note": "不生成、不修改 DOCX；答辩记录固定 v0.9 路径",
        "direction": DIRECTION,
        "students": [s for s, _ in STUDENTS],
        "metrics": metrics,
    }
    os.makedirs(OUT_DIR, exist_ok=True)
    path = os.path.join(OUT_DIR, "defense_diff_analysis_report.json")
    with io.open(path, "w", encoding="utf-8") as f:
        json.dump(report, f, ensure_ascii=False, indent=2)
    print("analysis-only saved:", path)
    print("score:", metrics["similarity_score"], "high:", metrics["high_similarity"])


if __name__ == "__main__":
    main()