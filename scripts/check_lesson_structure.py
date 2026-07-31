#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Validate a generated lesson-plan docx against the retained 39x6 template."""
import argparse
import re
import sys
from collections import Counter

from docx import Document
from docx.oxml.ns import qn


def grid(t):
    g = t._tbl.find(qn("w:tblGrid"))
    return [gc.get(qn("w:w")) for gc in g.findall(qn("w:gridCol"))]


def merge_sig(t):
    sig = []
    for row in t.rows:
        seen = []
        s = ""
        for c in row.cells:
            if c._tc in seen:
                s += "M"
            else:
                seen.append(c._tc)
                s += "U"
        sig.append(s)
    return sig


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("docx")
    ap.add_argument("--template", default=None)
    args = ap.parse_args()
    doc = Document(args.docx)
    t = doc.tables[0]
    checks = {"39x6": len(t.rows) == 39 and len(t.columns) == 6}
    if args.template:
        tpl = Document(args.template).tables[0]
        checks["grid"] = grid(t) == grid(tpl)
        checks["merges"] = merge_sig(t) == merge_sig(tpl)
    chain = re.findall(r"(\d+)\s*min", t.rows[7].cells[1].text)
    r10 = re.findall(r"（(\d+)分钟）", t.rows[10].cells[2].text)
    checks["90min"] = sum(map(int, chain)) == 90 if chain else False
    checks["R10=48"] = sum(int(x) for x in r10) == 48 if r10 else False
    colors = Counter()
    for row in t.rows:
        seen = []
        for c in row.cells:
            if c._tc in seen:
                continue
            seen.append(c._tc)
            for p in c.paragraphs:
                for r in p.runs:
                    rPr = r._r.find(qn("w:rPr"))
                    col = None
                    if rPr is not None:
                        e = rPr.find(qn("w:color"))
                        if e is not None:
                            col = e.get(qn("w:val"))
                    colors[col] += 1
    checks["red"] = colors.get("FF0000", 0) >= 1
    checks["blue"] = colors.get("2E75B6", 0) >= 1
    checks["purple"] = colors.get("7030A0", 0) >= 1
    checks["R38=6"] = len(t.rows[38].cells[1]._tc.findall(qn("w:p"))) == 6
    tail = doc.paragraphs[-1]._p
    pPr = tail.find(qn("w:pPr"))
    sp = pPr.find(qn("w:spacing")) if pPr is not None else None
    checks["tail0"] = sp is not None and sp.get(qn("w:line")) == "0"
    ok = True
    for k, v in checks.items():
        print(("PASS" if v else "FAIL"), k)
        ok = ok and v
    sys.exit(0 if ok else 1)


if __name__ == "__main__":
    main()
