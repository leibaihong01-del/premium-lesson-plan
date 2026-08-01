# -*- coding: utf-8 -*-
"""格式检测模块：模板结构、反思区、尾段。"""
import re

from docx import Document
from docx.oxml.ns import qn


def grid(t):
    g = t._tbl.find(qn("w:tblGrid"))
    return [gc.get(qn("w:w")) for gc in g.findall(qn("w:gridCol"))]


def merge_sig(t):
    sig = []
    for row in t.rows:
        seen = []
        s = ""
        for c in row.cells:
            if c._tc in seen:
                s += "M"
            else:
                seen.append(c._tc)
                s += "U"
        sig.append(s)
    return sig


def check_format(docx_path, template_path):
    doc = Document(docx_path)
    t = doc.tables[0]
    tpl = Document(template_path).tables[0]
    checks = {}
    checks["39x6"] = len(t.rows) == 39 and len(t.columns) == 6
    checks["grid"] = grid(t) == grid(tpl)
    checks["merges"] = merge_sig(t) == merge_sig(tpl)
    checks["R38_6"] = len(t.rows[38].cells[1]._tc.findall(qn("w:p"))) == 6
    tail = doc.paragraphs[-1]._p
    pPr = tail.find(qn("w:pPr"))
    sp = pPr.find(qn("w:spacing")) if pPr is not None else None
    checks["tail0"] = sp is not None and sp.get(qn("w:line")) == "0"
    passed = [k for k, v in checks.items() if v]
    score = round(20.0 * len(passed) / len(checks), 1)
    issues = [k for k, v in checks.items() if not v]
    return {"score": score, "checks": checks, "issues": issues}


def check_colors(docx_path):
    doc = Document(docx_path)
    colors = set()
    for t in doc.tables:
        seen = set()
        for row in t.rows:
            for c in row.cells:
                if c._tc in seen:
                    continue
                seen.add(c._tc)
                for p in c.paragraphs:
                    for r in p.runs:
                        rPr = r._r.find(qn("w:rPr"))
                        if rPr is not None:
                            e = rPr.find(qn("w:color"))
                            if e is not None:
                                colors.add(e.get(qn("w:val")))
    return colors


def check_title(docx_path):
    doc = Document(docx_path)
    return doc.paragraphs[0].text if doc.paragraphs else ""
