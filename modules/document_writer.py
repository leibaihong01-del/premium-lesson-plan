# -*- coding: utf-8 -*-
"""文档生成模块：模板复制 + run 级填充 + 格式规则。"""
import copy
import os
import re
import shutil

from docx import Document
from docx.oxml.ns import qn


GRAY = "1F1F1F"
RED = "FF0000"
BLUE = "2E75B6"
PURPLE = "7030A0"


def set_text(r, text):
    for t in r.findall(qn("w:t")):
        r.remove(t)
    t = r.makeelement(qn("w:t"), {})
    t.text = text
    if text != text.strip():
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)


def set_cell_lines(cell, lines):
    tc = cell._tc
    ps = tc.findall(qn("w:p"))
    if not ps:
        return
    n = len(lines)
    while len(ps) < n:
        newp = copy.deepcopy(ps[-1])
        ps[-1].addnext(newp)
        ps = tc.findall(qn("w:p"))
    for extra in ps[n:]:
        tc.remove(extra)
    ps = tc.findall(qn("w:p"))
    for i, p in enumerate(ps):
        runs = p.findall(qn("w:r"))
        if not runs:
            continue
        r = runs[0]
        set_text(r, lines[i])
        for extra in runs[1:]:
            p.remove(extra)


def style_run(r, color, bold):
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = r.makeelement(qn("w:rPr"), {})
        r.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    for a in ("ascii", "hAnsi", "eastAsia", "cs"):
        rFonts.set(qn("w:" + a), "楷体")
    b = rPr.find(qn("w:b"))
    if bold:
        if b is None:
            b = rPr.makeelement(qn("w:b"), {})
            rPr.append(b)
    else:
        if b is not None:
            rPr.remove(b)
    bCs = rPr.find(qn("w:bCs"))
    if bCs is not None:
        rPr.remove(bCs)
    c = rPr.find(qn("w:color"))
    if c is None:
        c = rPr.makeelement(qn("w:color"), {})
        rPr.append(c)
    c.set(qn("w:val"), color)
    if c.get(qn("w:themeColor")):
        del c.attrib[qn("w:themeColor")]
    for tag in ("w:sz", "w:szCs"):
        e = rPr.find(qn(tag))
        if e is None:
            e = rPr.makeelement(qn(tag), {})
            rPr.append(e)
        e.set(qn("w:val"), "21")


def classify(text):
    if text.startswith("【思政融入】"):
        return RED, True
    if text.startswith("【随堂练习】") or text.startswith("【实操任务】"):
        return BLUE, False
    if re.match(r"^任务\d+", text):
        return PURPLE, True
    if text.startswith("知识点"):
        return PURPLE, False
    if text.startswith("教学重点：") or text.startswith("教学难点："):
        return GRAY, True
    if text.startswith("标题：") or re.match(r"^[一二三四五六七八九十]+、", text):
        return GRAY, True
    return GRAY, False


def format_paragraph(p):
    text = p.text
    color, bold = classify(text)
    runs = p._p.findall(qn("w:r"))
    for extra in runs[1:]:
        p._p.remove(extra)
    if not runs:
        return
    r = runs[0]
    style_run(r, color, bold)
    set_text(r, text)
    if color == RED and text.startswith("【思政融入】"):
        label = "【思政融入】"
        set_text(r, label)
        r2 = copy.deepcopy(r)
        style_run(r2, RED, False)
        set_text(r2, text[len(label):])
        r.addnext(r2)


def rebuild_mixed(p):
    text = "".join(t.text or "" for t in p._p.iter(qn("w:t")))
    parts = re.split(r"(课前|课中|课后)", text)
    runs = p._p.findall(qn("w:r"))
    base = runs[0] if runs else p._p.makeelement(qn("w:r"), {})
    for r in runs:
        p._p.remove(r)
    for seg in parts:
        if not seg:
            continue
        bold = seg in ("课前", "课中", "课后")
        r = copy.deepcopy(base)
        style_run(r, PURPLE if bold else GRAY, bold)
        set_text(r, seg)
        p._p.append(r)


def split_long(cell, limit=55):
    ps = cell._tc.findall(qn("w:p"))
    i = 0
    while i < len(ps):
        p = ps[i]
        text = "".join(t.text or "" for t in p.iter(qn("w:t")))
        skip = text.startswith(("【思政融入】", "【随堂练习】", "【实操任务】", "任务", "知识点", "“"))
        if len(text) > limit and not skip:
            cut = None
            for marker in ("；", "，", ",", "。"):
                pos = text.find(marker)
                if 35 <= pos <= 70:
                    cut = pos + 1
                    break
            if cut is None:
                pos = text.find("，", 40)
                if pos > 0:
                    cut = pos + 1
            if cut:
                part1, part2 = text[:cut], text[cut:]
                runs = p.findall(qn("w:r"))
                if runs:
                    for extra in runs[1:]:
                        p.remove(extra)
                    set_text(runs[0], part1)
                    newp = copy.deepcopy(p)
                    nruns = newp.findall(qn("w:r"))
                    for extra in nruns[1:]:
                        newp.remove(extra)
                    if nruns:
                        set_text(nruns[0], part2)
                    p.addnext(newp)
                    ps.insert(i + 1, newp)
        i += 1


def fix_tail(doc):
    tail = doc.paragraphs[-1]._p
    pPr = tail.find(qn("w:pPr"))
    if pPr is None:
        pPr = tail.makeelement(qn("w:pPr"), {})
        tail.insert(0, pPr)
    sp = pPr.find(qn("w:spacing"))
    if sp is None:
        sp = pPr.makeelement(qn("w:spacing"), {})
        pPr.insert(0, sp)
    sp.set(qn("w:line"), "0")
    sp.set(qn("w:lineRule"), "exact")
    sp.set(qn("w:before"), "0")
    sp.set(qn("w:after"), "0")
    rPr = pPr.find(qn("w:rPr"))
    if rPr is None:
        rPr = pPr.makeelement(qn("w:rPr"), {})
        pPr.append(rPr)
    for tag in ("w:sz", "w:szCs"):
        e = rPr.find(qn(tag))
        if e is None:
            e = rPr.makeelement(qn(tag), {})
            rPr.append(e)
        e.set(qn("w:val"), "1")


def apply_row_heights(tbl):
    for ri in (3, 8, 24):
        tr = tbl.rows[ri]._tr
        trPr = tr.find(qn("w:trPr"))
        if trPr is not None:
            h = trPr.find(qn("w:trHeight"))
            if h is not None:
                trPr.remove(h)
    for ri, h in [(10, 500), (11, 500), (12, 500), (13, 500), (17, 500), (30, 400), (31, 500), (32, 500)]:
        tr = tbl.rows[ri]._tr
        trPr = tr.find(qn("w:trPr"))
        if trPr is None:
            trPr = tr.makeelement(qn("w:trPr"), {})
            tr.insert(0, trPr)
        e = trPr.find(qn("w:trHeight"))
        if e is None:
            e = trPr.makeelement(qn("w:trHeight"), {})
            trPr.append(e)
        e.set(qn("w:val"), str(h))
        e.set(qn("w:hRule"), "atLeast")
    for ri, h in [(35, 400), (36, 400), (37, 406)]:
        tr = tbl.rows[ri]._tr
        trPr = tr.find(qn("w:trPr"))
        if trPr is None:
            trPr = tr.makeelement(qn("w:trPr"), {})
            tr.insert(0, trPr)
        e = trPr.find(qn("w:trHeight"))
        if e is None:
            e = trPr.makeelement(qn("w:trHeight"), {})
            trPr.append(e)
        e.set(qn("w:val"), str(h))
        e.set(qn("w:hRule"), "atLeast")


def generate_document(template_path, output_path, fills, project_kind="lesson"):
    shutil.copy2(template_path, output_path)
    doc = Document(output_path)
    tbl = doc.tables[0]
    for row, col, lines in fills:
        set_cell_lines(tbl.rows[row].cells[col], lines)
    apply_row_heights(tbl)
    targets = []
    for ri in [3, 4, 5, 6, 7, 8]:
        targets.append((ri, [1]))
    for ri in [10, 11, 12, 13]:
        targets.append((ri, [2]))
    targets.append((17, [1, 5]))
    for ri in range(20, 28):
        targets.append((ri, [1, 5]))
    for ri in range(30, 33):
        targets.append((ri, [1, 5]))
    for ri, cols in targets:
        for ci in cols:
            for p in tbl.rows[ri].cells[ci].paragraphs:
                format_paragraph(p)
    for p in tbl.rows[7].cells[1].paragraphs:
        rebuild_mixed(p)
    split_long(tbl.rows[24].cells[1])
    if project_kind == "practice":
        for ri in (38,):
            tr = tbl.rows[ri]._tr
            trPr = tr.find(qn("w:trPr"))
            if trPr is None:
                trPr = tr.makeelement(qn("w:trPr"), {})
                tr.insert(0, trPr)
            if trPr.find(qn("w:cantSplit")) is None:
                trPr.append(trPr.makeelement(qn("w:cantSplit"), {}))
    fix_tail(doc)
    doc.save(output_path)
    return output_path
