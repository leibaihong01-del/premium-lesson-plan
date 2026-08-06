# -*- coding: utf-8 -*-
"""内容与教学专业度检测模块。"""
import re

from docx import Document


FORBIDDEN = ["信号设备检修", "联锁故障处理", "CBTC参数配置", "设备系统集成调试", "车辆检修规程", "故障诊断方法"]


def cell_text(t, row, col):
    return t.rows[row].cells[col].text


def safe_cell(t, row, col):
    """行号越界时返回空字符串，兼容不同模板行列结构。"""
    if row >= len(t.rows) or col >= len(t.columns):
        return ""
    return cell_text(t, row, col)


def check_content(docx_path):
    doc = Document(docx_path)
    t = doc.tables[0]
    checks = {}
    checks["goals"] = bool(safe_cell(t, 3, 1).strip())
    checks["keypoints"] = bool(safe_cell(t, 4, 1).strip())
    chain = re.findall(r"(\d+)\s*min", safe_cell(t, 7, 1))
    checks["90min"] = sum(map(int, chain)) == 90 if chain else False
    r10 = re.findall(r"（(\d+)分钟）", safe_cell(t, 10, 2))
    checks["R10_48"] = sum(int(x) for x in r10) == 48 if r10 else False
    checks["process"] = all(bool(safe_cell(t, ri, 1).strip()) for ri in range(20, 28))
    checks["after"] = all(bool(safe_cell(t, ri, 1).strip()) for ri in (30, 31, 32))
    passed = [k for k, v in checks.items() if v]
    score = round(25.0 * len(passed) / len(checks), 1)
    return {"score": score, "checks": checks, "issues": [k for k, v in checks.items() if not v]}


def check_teaching(docx_path):
    doc = Document(docx_path)
    t = doc.tables[0]
    text = "\n".join(c.text for row in t.rows for c in row.cells)
    checks = {}
    checks["思政"] = ("思政育人目标" in text) and ("【思政融入】" in text)
    checks["随堂/实操"] = ("【随堂练习】" in text) or ("【实操任务】" in text)
    checks["分层/实训任务"] = ("基础作业" in text) or ("实训报告" in text)
    checks["岗位情境"] = ("【岗位情境】" in text) or ("岗位" in text)
    checks["无越界"] = not any(w in text for w in FORBIDDEN)
    passed = [k for k, v in checks.items() if v]
    score = round(25.0 * len(passed) / len(checks), 1)
    return {"score": score, "checks": checks, "issues": [k for k, v in checks.items() if not v]}


def check_format_style(docx_path):
    from .format_checker import check_colors, check_title

    colors = check_colors(docx_path)
    title = check_title(docx_path)
    checks = {
        "red": "FF0000" in colors,
        "blue": "2E75B6" in colors,
        "purple": "7030A0" in colors,
        "title": title.strip().endswith("课程教案"),
    }
    passed = [k for k, v in checks.items() if v]
    score = round(20.0 * len(passed) / len(checks), 1)
    return {"score": score, "checks": checks, "issues": [k for k, v in checks.items() if not v]}