# -*- coding: utf-8 -*-
"""模板解析模块：提取页面、表格、页眉页脚结构。"""
from docx import Document
from docx.oxml.ns import qn


def grid(t):
    g = t._tbl.find(qn("w:tblGrid"))
    return [gc.get(qn("w:w")) for gc in g.findall(qn("w:gridCol"))]


def merge_sig(t):
    sig = []
    for row in t.rows:
        seen = []
        s = ""
        for c in row.cells:
            if c._tc in seen:
                s += "M"
            else:
                seen.append(c._tc)
                s += "U"
        sig.append(s)
    return sig


def parse_template(path):
    doc = Document(path)
    sec = doc.sections[0]
    info = {
        "path": path,
        "page": {
            "width_emu": sec.page_width,
            "height_emu": sec.page_height,
            "margins": {
                "top": sec.top_margin,
                "right": sec.right_margin,
                "bottom": sec.bottom_margin,
                "left": sec.left_margin,
            },
        },
        "tables": [],
        "header_parts": [],
        "footer_parts": [],
        "first_title": doc.paragraphs[0].text if doc.paragraphs else "",
    }
    for ti, t in enumerate(doc.tables):
        info["tables"].append({
            "index": ti,
            "rows": len(t.rows),
            "cols": len(t.columns),
            "grid": grid(t),
            "merge_sig": merge_sig(t),
        })
    rels = doc.part.rels
    for rid, rel in rels.items():
        if "header" in rel.reltype:
            info["header_parts"].append(rel.target_part.partname if rel.target_part else rid)
        if "footer" in rel.reltype:
            info["footer_parts"].append(rel.target_part.partname if rel.target_part else rid)
    return info
