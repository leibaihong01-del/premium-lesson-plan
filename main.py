# -*- coding: utf-8 -*-
"""CourseAgent 总控制器：规划→生成→检测→修复→打包→复盘→记忆。"""
import argparse
import json
import os
import shutil
import sys

import yaml
from docx import Document
from docx.oxml.ns import qn

from agents import learner, planner, reviewer, router_agent, writer
from modules import packager, runtime, template_parser, visual_checker


ROOT = os.path.dirname(os.path.abspath(__file__))
CONFIG = os.path.join(ROOT, "config", "agent_rules.yaml")
INPUT = os.path.join(ROOT, "input")
OUTPUT = os.path.join(ROOT, "output")
REPORTS = os.path.join(ROOT, "reports")
MEMORY = os.path.join(ROOT, "memory")


def load_config():
    with open(CONFIG, encoding="utf-8") as f:
        return yaml.safe_load(f)


def fix_tail_spacing(path):
    doc = Document(path)
    tail = doc.paragraphs[-1]._p
    pPr = tail.find(qn("w:pPr"))
    if pPr is None:
        pPr = tail.makeelement(qn("w:pPr"), {})
        tail.insert(0, pPr)
    sp = pPr.find(qn("w:spacing"))
    if sp is None:
        sp = pPr.makeelement(qn("w:spacing"), {})
        pPr.insert(0, sp)
    sp.set(qn("w:line"), "0")
    sp.set(qn("w:lineRule"), "exact")
    doc.save(path)


def fix_title(path, title):
    doc = Document(path)
    p = doc.paragraphs[0]
    if p.runs:
        full = "".join(r.text or "" for r in p.runs)
        if title and title not in full:
            p.runs[0].text = title
            for extra in p.runs[1:]:
                extra.text = ""
            doc.save(path)


def repair(path, kind, title):
    actions = runtime.load_cache(os.path.join(MEMORY, "repair_actions.json"))
    fix_tail_spacing(path)
    actions.setdefault("tail_spacing", {"issue": "tail0", "action": "fix_tail_spacing", "count": 0})["count"] += 1
    if title:
        fix_title(path, title)
        actions.setdefault("title_fix", {"issue": "title", "action": "fix_title", "count": 0})["count"] += 1
    if kind == "practice":
        doc = Document(path)
        tr = doc.tables[0].rows[38]._tr
        trPr = tr.find(qn("w:trPr"))
        if trPr is None:
            trPr = tr.makeelement(qn("w:trPr"), {})
            tr.insert(0, trPr)
        if trPr.find(qn("w:cantSplit")) is None:
            trPr.append(trPr.makeelement(qn("w:cantSplit"), {}))
        doc.save(path)
        actions.setdefault("reflection_no_split", {"issue": "practice_reflection", "action": "cantSplit_row38", "count": 0})["count"] += 1
    runtime.save_cache(os.path.join(MEMORY, "repair_actions.json"), actions)


def run(args):
    config = load_config()
    min_score = config["quality"]["minimum_score"]
    max_loops = config["repair"]["max_loops"]
    weights = config["quality"]["weights"]
    os.makedirs(OUTPUT, exist_ok=True)
    os.makedirs(REPORTS, exist_ok=True)
    os.makedirs(INPUT, exist_ok=True)

    plan, _ = planner.analyze_task(
        config, args.task, args.template, args.existing or args.fills or "", REPORTS
    )
    tpl_info = template_parser.parse_template(args.template)
    with open(os.path.join(REPORTS, "template_structure.json"), "w", encoding="utf-8") as f:
        json.dump(tpl_info, f, ensure_ascii=False, indent=2)

    file_meta = {
        "task_type": args.kind if args.kind in ("lesson", "practice") else "docx",
        "domain": "education",
        "pages": 0,
        "tables": tpl_info["tables"][0]["rows"] if tpl_info.get("tables") else 0,
        "images": len(tpl_info.get("header_parts", [])) + len(tpl_info.get("footer_parts", [])),
        "file_count": 1 if args.existing else 2,
        "complex_layout": bool(tpl_info.get("tables")),
    }
    route = router_agent.route(file_meta, args.task, force_profile=args.profile)
    level = route.get("recommended_compute", "medium")
    print("路由决策: 复杂度=%s 置信度=%s 等级=%s 来源=%s 原因=%s" % (
        route.get("complexity"), route.get("confidence"), level,
        route.get("source"), " / ".join(route.get("reason", []))))

    work = os.path.join(OUTPUT, "_work.docx")
    if args.fills:
        writer.write_document(args.template, work, fills_path=args.fills, kind=args.kind)
        target = work
    elif args.existing:
        shutil.copy2(args.existing, work)
        target = work
    else:
        raise SystemExit("请提供 --existing 或 --fills")

    cache_path = os.path.join(REPORTS, "review_cache.json")
    cache_key_src = args.existing or (args.fills or "")
    cached, reused = runtime.cached_review(cache_key_src, args.template, cache_path) if cache_key_src else (None, False)
    report = None
    loops = 0
    if reused:
        report = runtime.normalize_cached(cached)
        report["file"] = args.project
        print("缓存命中，跳过重复检测")
    else:
        loops_limit = config["compute"]["profiles"][level].get("loops", 3)
        while loops < min(loops_limit, max_loops):
            report = reviewer.review(target, args.template, REPORTS, weights)
            report["loops"] = loops + 1
            report["file"] = args.project
            if report["final"] == "PASS":
                break
            repair(target, args.kind, args.title)
            loops += 1
            report["rules_updated"] = [
                {"id": "R-" + str(loops), "rule": "检测未达标自动修复并复检", "source": "main.repair"}
            ]
        report["loops"] = loops + 1
    if level == "high" and args.pdf and os.path.exists(args.pdf):
        report["visual"] = visual_checker.check_pdf(args.pdf)
        print("视觉检查: 页数=%s 结论=%s 问题=%s" % (
            report["visual"].get("pages"), "通过" if report["visual"]["ok"] else "未通过",
            report["visual"].get("issues")))
    router_agent.record(file_meta, report["total_score"], report["final"] == "PASS",
                        report["loops"], level)
    runtime.put_cache(cache_path, cache_key_src or target, args.template, report)

    final_docx = os.path.join(OUTPUT, args.project + "_优化后版本.docx")
    shutil.copy2(target, final_docx)
    base, dirs = packager.package(args.project, final_docx, REPORTS, MEMORY, OUTPUT)
    reflection = learner.learn(report, MEMORY, REPORTS)

    print("=" * 50)
    print("项目:", args.project)
    print("综合质量:", report["total_score"], "| 结论:", report["final"], "| 循环:", report["loops"])
    print("模板:", report["template_score"], "内容:", report["content_score"],
          "教学:", report["teaching_score"], "格式:", report["format_score"])
    print("输出目录:", base)
    print("复盘日志:", reflection)


def main():
    ap = argparse.ArgumentParser(description="CourseAgent 课程文档智能体")
    ap.add_argument("--template", required=True, help="模板 docx")
    ap.add_argument("--existing", help="已生成 docx（评审/打包模式）")
    ap.add_argument("--fills", help="fills.json 路径（生成模式）")
    ap.add_argument("--project", default="课程文档", help="项目名称")
    ap.add_argument("--kind", choices=["lesson", "practice"], default="lesson")
    ap.add_argument("--title", default="", help="期望正文标题")
    ap.add_argument("--task", default="根据模板生成精品课程教案并完成质量闭环", help="任务描述")
    ap.add_argument("--profile", choices=["auto", "low", "medium", "high"], default="auto", help="算力档位")
    ap.add_argument("--requirement", choices=["normal", "formal", "excellent"], default="excellent", help="质量要求")
    ap.add_argument("--pdf", default="", help="高算力档视觉检查用PDF路径")
    args = ap.parse_args()
    run(args)


if __name__ == "__main__":
    main()
