# -*- coding: utf-8 -*-
"""精品课程建设专家：五维诊断与提升建议。"""
import re

from docx import Document


WEIGHTS = {"教学逻辑": 25, "内容体系": 25, "职业特色": 20, "创新设计": 15, "评价体系": 15}


def _text(docx_path):
    doc = Document(docx_path)
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        seen = set()
        for row in t.rows:
            for c in row.cells:
                if c._tc in seen:
                    continue
                seen.add(c._tc)
                parts.append(c.text)
    return "\n".join(parts)


def analyze(docx_path):
    text = _text(docx_path)
    doc = Document(docx_path)
    tbl = doc.tables[0]
    r7 = tbl.rows[7].cells[1].text
    r10 = tbl.rows[10].cells[2].text
    dims = {}
    # 教学逻辑
    chain = re.findall(r"(\d+)\s*min", r7)
    r10_times = re.findall(r"（(\d+)分钟）", r10)
    logic = 0
    logic += 30 if sum(map(int, chain)) == 90 else 0
    logic += 30 if sum(int(x) for x in r10_times) == 48 else 0
    logic += 20 if "任务实施" in text and "任务总结" in text else 0
    logic += 20 if "课前" in text and "课中" in text and "课后" in text else 0
    dims["教学逻辑"] = logic
    # 内容体系
    content = 0
    content += 20 if "教学目标" in text and "知识目标" in text and "能力目标" in text else 0
    content += 20 if "教学重" in text and "教学难" in text else 0
    content += 20 if "板书" in text and "标题：" in text else 0
    content += 20 if "基础作业" in text or "实训报告" in text else 0
    content += 20 if "教学反思" in text and "改进方法" in text else 0
    dims["内容体系"] = content
    # 职业特色
    prof = 0
    prof += 40 if "【岗位情境】" in text or "岗位" in text else 0
    prof += 30 if "实操" in text or "实训" in text else 0
    prof += 30 if "岗位育人落点" in text else 0
    dims["职业特色"] = prof
    # 创新设计
    innov = 0
    innov += 30 if any(k in text for k in ("案例", "视频", "动画", "仿真")) else 0
    innov += 30 if "任务" in text and "知识点" in text else 0
    innov += 40 if "【思政融入】" in text else 0
    dims["创新设计"] = innov
    # 评价体系
    eval_ = 0
    eval_ += 30 if "评价" in text and ("评分" in text or "维度" in text) else 0
    eval_ += 20 if "【随堂练习】" in text or "【实操任务】" in text else 0
    eval_ += 20 if "学习通" in text else 0
    eval_ += 30 if "过程" in text or "多元评价" in text else 0
    dims["评价体系"] = eval_
    total = round(sum(dims[k] * WEIGHTS[k] for k in dims) / 100, 1)
    level = "标志性成果" if total >= 95 else "精品课程" if total >= 90 else "优质课程" if total >= 80 else "普通课程"
    gaps = {k: 100 - v for k, v in dims.items() if v < 90}
    suggestions = {k: _suggest(k) for k in sorted(gaps, key=gaps.get, reverse=True)}
    return {"file": docx_path, "dimensions": dims, "total": total, "level": level, "gaps": gaps, "suggestions": suggestions}


def _suggest(dim):
    return {
        "教学逻辑": "强化90分钟时间链与任务时间分配，确保教学做一体48分钟闭合",
        "内容体系": "补全目标层级、板书结构、分层作业与反思区内容",
        "职业特色": "增加岗位情境、实操任务与岗位育人落点",
        "创新设计": "引入行业案例、项目化任务与思政融入点",
        "评价体系": "完善多元评价维度、随堂/实操任务与学习通测评",
    }[dim]
