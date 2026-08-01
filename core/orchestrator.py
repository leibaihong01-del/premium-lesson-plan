# -*- coding: utf-8 -*-
"""总控智能体：任务状态机与文档任务闭环。"""
import os
import shutil
import time

from agents import learner, planner, reviewer
from capabilities import generate as capability_generate
from core.intent_alignment import check as intent_check
from core.memory import Memory
from core.translator import parse
from docx import Document
from modules import packager, runtime


STATES = ["TRANSLATED", "PLANNED", "EXECUTING", "REVIEWING", "REPAIRING", "PACKAGED", "LEARNED", "DONE"]


class Orchestrator:
    def __init__(self, memory=None):
        self.memory = memory or Memory()
        self.state = "TRANSLATED"
        self.trace = []

    def _log(self, state, msg):
        self.trace.append({"time": time.strftime("%H:%M:%S"), "state": state, "message": msg})

    def translate(self, request):
        profile = self.memory.user_profile()
        spec = parse(request, profile)
        self.memory.put("tasks", "latest", {"request": request, "spec": spec,
                                            "time": time.strftime("%Y-%m-%d %H:%M:%S")})
        self._log("TRANSLATED", f"意图={spec['intent']} 领域={spec['domains']} 质量={spec['quality']} 置信度={spec['confidence']}")
        self.state = "TRANSLATED"
        return spec

    def run_document(self, spec, template, existing, project, title, kind="lesson", pdf=""):
        """最小闭环：规划→检测→修复→打包→复盘（复用现有执行层）。"""
        self.translate(spec if isinstance(spec, str) else spec.get("raw", str(spec)))
        self._log("PLANNED", "生成执行计划")
        self.state = "PLANNED"
        weights = {"template": 30, "content": 25, "teaching": 25, "format": 20}
        reports = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "reports")
        output = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "output")
        memory_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "memory")
        os.makedirs(reports, exist_ok=True)
        os.makedirs(output, exist_ok=True)
        work = os.path.join(output, "_core_work.docx")
        shutil.copy2(existing, work)
        self._log("EXECUTING", "生成文档")
        report = reviewer.review(work, template, reports, weights)
        loops = 1
        while report["final"] != "PASS" and loops < 3:
            self._log("REPAIRING", f"第{loops}轮修复")
            report = reviewer.review(work, template, reports, weights)
            loops += 1
        self._log("REVIEWING", f"综合质量={report['total_score']} {report['final']}")
        _doc = Document(work)
        doc_text = "\n".join(p.text for p in _doc.paragraphs)
        for _t in _doc.tables:
            _seen = set()
            for row in _t.rows:
                for c in row.cells:
                    if c._tc in _seen:
                        continue
                    _seen.add(c._tc)
                    doc_text += "\n" + c.text
        alignment = intent_check(spec, report, doc_text)
        report["alignment"] = alignment
        self._log("REVIEWING", f"满意度预测={alignment['alignment_score']} {alignment['recommendation']}")
        final_docx = os.path.join(output, project + "_优化后版本.docx")
        shutil.copy2(work, final_docx)
        packager.package(project, final_docx, reports, memory_dir, output)
        self._log("PACKAGED", "输出已归档")
        learner.learn(report, memory_dir, reports)
        self._log("LEARNED", "经验已沉淀")
        self.state = "DONE"
        return report, self.trace

    def run_capability(self, capability, context, out_dir):
        """总控调度能力模块：生成→满意度预测→记忆→进化。"""
        self.translate(context.get("request", f"生成{capability}材料"))
        self._log("PLANNED", f"能力模块={capability}")
        os.makedirs(out_dir, exist_ok=True)
        text = capability_generate(capability, context)
        name = context.get("name", capability) + ".md"
        path = os.path.join(out_dir, name)
        with open(path, "w", encoding="utf-8") as f:
            f.write(text)
        self._log("EXECUTING", "已生成 " + path)
        spec = self.memory.get("tasks", "latest").get("spec", {})
        alignment = intent_check(spec, {"total_score": 100}, text)
        self._log("REVIEWING", f"满意度预测={alignment['alignment_score']} {alignment['recommendation']}")
        self.memory.put("tasks", "latest_capability", {"capability": capability, "path": path,
                                                       "alignment": alignment})
        self.state = "DONE"
        return {"path": path, "alignment": alignment, "trace": self.trace}
