# -*- coding: utf-8 -*-
"""Defense record layout normalizer.

Template-derived rules from the 3-version defense layout analysis:
- table row height: restore template 12382 twips, rule "atLeast"
- first 7 paragraphs: exact line spacing 22pt
- remaining paragraphs: exact line spacing 18pt
- abnormal empty paragraphs are cleaned
- cell text: 12pt, 宋体 for CJK, Times New Roman for Latin/digits
"""
import os
import re

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ENABLED = os.environ.get("DEFENSE_LAYOUT_NORMALIZE", "0") == "1"


def _clean_paragraph(p):
    """Remove empty trailing paragraphs except the last one in a cell/body."""
    text = "".join(r.text for r in p.runs).strip()
    if text:
        return False
    for child in list(p._p):
        tag = child.tag
        if tag != qn("w:pPr"):
            p._p.remove(child)
    return True


def _set_exact_spacing(p, pts):
    p_pr = p._p.get_or_add_pPr()
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        p_pr.append(spacing)
    spacing.set(qn("w:line"), str(int(pts * 20)))
    spacing.set(qn("w:lineRule"), "exact")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")


def _fix_run_font(run):
    run.font.size = Pt(12)
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "宋体")


def _restore_table_rows(doc):
    for table in doc.tables:
        for row in table.rows:
            tr_pr = row._tr.get_or_add_trPr()
            tr_height = tr_pr.find(qn("w:trHeight"))
            if tr_height is None:
                tr_height = OxmlElement("w:trHeight")
                tr_pr.append(tr_height)
            tr_height.set(qn("w:val"), "12382")
            tr_height.set(qn("w:hRule"), "atLeast")


def _normalize_container(paragraphs, first_7=22, rest=18):
    kept = 0
    for p in paragraphs:
        if _clean_paragraph(p):
            continue
        _set_exact_spacing(p, first_7 if kept < 7 else rest)
        for run in p.runs:
            _fix_run_font(run)
        kept += 1


def normalize(docx_path):
    if not ENABLED or not os.path.isfile(docx_path):
        return False
    doc = Document(docx_path)
    _normalize_container(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            seen = set()
            for cell in row.cells:
                if cell._tc in seen:
                    continue
                seen.add(cell._tc)
                _normalize_container(cell.paragraphs)
    _restore_table_rows(doc)
    doc.save(docx_path)
    return True


def main():
    import sys
    if len(sys.argv) != 2:
        print("usage: defense_layout_normalizer.py <docx>")
        return 1
    print("normalized" if normalize(sys.argv[1]) else "skipped")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
