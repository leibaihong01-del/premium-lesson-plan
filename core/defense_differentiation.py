# -*- coding: utf-8 -*-
"""答辩记录表同方向差异化控制（v0.1）。

仅对同方向批量生成时生效：
- 相似度检测：首句重复 / 高频句式 / 关键词重复率；
- 高度相似时基于真实材料重组表达（只重排已有句子，不虚构技术内容）；
- 单个学生生成不改变内部逻辑。
"""
import copy
import hashlib
import io
import json
import os
import re

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


DEFAULT_PROFILE_PATH = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "skills", "graduation_design", "defense_record", "rules", "direction_profiles.json")


def load_profiles(path=DEFAULT_PROFILE_PATH):
    with io.open(path, encoding="utf-8") as f:
        return json.load(f)


def normalize(s):
    return re.sub(r"\s+", "", s or "").replace("\u3000", "")


def split_sentences(text):
    parts = re.split(r"(?<=[。！？；\n])", text or "")
    return [p.strip() for p in parts if p.strip()]


def _first_sentences(texts):
    out = []
    for t in texts:
        sents = split_sentences(t)
        out.append(normalize(sents[0]) if sents else "")
    return out


def _ngram(text, n=4):
    return [text[i:i + n] for i in range(max(0, len(text) - n + 1))]


def analyze_similarity(texts, direction=None, profiles=None):
    """返回同方向文本相似度指标。"""
    profiles = profiles or load_profiles()
    directions = profiles.get("directions", profiles)
    n = len(texts)
    firsts = _first_sentences(texts)
    dup_first = 0
    for i in range(n):
        for j in range(i + 1, n):
            if firsts[i] and firsts[i] == firsts[j]:
                dup_first += 1
    max_pairs = max(1, n * (n - 1) // 2)
    first_dup_rate = dup_first / max_pairs

    pattern_counts = {}
    for t in texts:
        seen = set(_ngram(normalize(t)))
        for g in seen:
            pattern_counts[g] = pattern_counts.get(g, 0) + 1
    repeated = {g: c for g, c in pattern_counts.items() if c >= 2}
    pattern_rate = min(1.0, len(repeated) / max(1, len(pattern_counts)))

    keywords = []
    if direction and direction in directions:
        keywords = directions.get(direction, {}).get("keywords", [])
    kw_hits = {k: sum(1 for t in texts if k in t) for k in keywords}
    kw_rate = 0.0
    if keywords:
        kw_rate = sum(1 for k, c in kw_hits.items() if c >= 2) / len(keywords)

    top_patterns = sorted(repeated.items(), key=lambda kv: (-kv[1], kv[0]))[:5]
    templates = []
    for t in texts:
        tpl = []
        for s in split_sentences(t):
            ns = normalize(s)
            if ns:
                tpl.append(ns[:6])
        templates.append(tpl)
    tpl_counts = {}
    for tpl in templates:
        for tm in set(tpl):
            tpl_counts[tm] = tpl_counts.get(tm, 0) + 1
    repeated_tpl = {tm: c for tm, c in tpl_counts.items() if c >= 2}
    template_dup_rate = min(1.0, len(repeated_tpl) / max(1, len(tpl_counts)))
    score = (first_dup_rate * 0.45 + template_dup_rate * 0.40 + kw_rate * 0.15)
    return {
        "count": n,
        "first_dup_rate": round(first_dup_rate, 3),
        "template_dup_rate": round(template_dup_rate, 3),
        "pattern_rate": round(pattern_rate, 3),
        "keyword_overlap_rate": round(kw_rate, 3),
        "similarity_score": round(score, 3),
        "high_similarity": score >= 0.5,
        "repeated_patterns": [g for g, c in top_patterns],
        "keyword_hits": kw_hits,
    }


def _group_sentences(sentences):
    """按序号标记分组；无标记时按自然句组。"""
    groups = []
    cur = []
    marker = re.compile(r"^[（(]?[一二三四五六七八九十\d]+[）)、.．]")
    for s in sentences:
        if marker.match(s) and cur:
            groups.append(cur)
            cur = [s]
        else:
            cur.append(s)
    if cur:
        groups.append(cur)
    if len(groups) <= 1 and len(sentences) >= 4:
        groups = [sentences[i:i + 2] for i in range(0, len(sentences), 2)]
    return groups


def rewrite_answer(answer, ctx, profile):
    """基于真实材料重组表达，只重排已有句子，不新增技术内容。"""
    sentences = split_sentences(answer)
    if len(sentences) < 4:
        return {"text": answer, "changed": False, "reason": "句子过少，保持原文"}
    opening = sentences[:1]
    closing = sentences[-1:]
    body = sentences[1:-1]
    groups = _group_sentences(body)
    if len(groups) < 2:
        return {"text": answer, "changed": False, "reason": "结构单一，保持原文"}

    dims = profile.get("expression_dimensions", [])
    seed = (ctx.get("student", "") + ctx.get("topic", "")).encode("utf-8")
    variant = int(hashlib.md5(seed).hexdigest(), 16) % max(1, len(dims))
    if variant % 2 == 0:
        groups = groups[::-1]
    else:
        groups = groups[1:] + groups[:1]

    parts = [g for g in groups]
    new_text = "".join(opening + [s for g in parts for s in g] + closing)
    if normalize(new_text) == normalize(answer):
        return {"text": answer, "changed": False, "reason": "重排后无变化"}
    return {
        "text": new_text,
        "changed": True,
        "reason": "同方向高相似，基于真实材料重组表达顺序",
        "variant": dims[variant] if dims else "",
    }


def process_batch(items, profiles=None):
    """items: [{student, topic, direction, answer}]，同方向批量差异化。"""
    profiles = profiles or load_profiles()
    by_direction = {}
    for it in items:
        by_direction.setdefault(it.get("direction", ""), []).append(it)
    report = {"metrics": {}, "rewrites": {}, "kept": []}
    for direction, group in by_direction.items():
        texts = [it.get("answer", "") for it in group]
        metrics = analyze_similarity(texts, direction=direction, profiles=profiles)
        report["metrics"][direction] = metrics
        if metrics["high_similarity"]:
            for it in group:
                directions = profiles.get("directions", profiles)
                profile = directions.get(direction, {})
                ctx = {
                    "student": it.get("student", ""),
                    "topic": it.get("topic", ""),
                    "station": it.get("station", ""),
                }
                rw = rewrite_answer_structured(it.get("answer", ""), ctx, profile)
                report["rewrites"][it.get("student", "")] = rw
        else:
            for it in group:
                report["kept"].append(it.get("student", ""))
    return report



def get_professional_profile(professional, profiles=None):
    profiles = profiles or load_profiles()
    return profiles.get("professionals", {}).get(professional, {})


def extract_keywords(text, top=12):
    import re as _re
    from collections import Counter
    pure = _re.sub(r"[^\u4e00-\u9fffA-Za-z0-9]", "", text or "")
    grams = []
    cjk = [ch for ch in pure if "\u4e00" <= ch <= "\u9fff"]
    for i in range(len(cjk) - 1):
        grams.append(cjk[i] + cjk[i + 1])
    stop = {"??", "??", "??", "??", "??", "??", "??", "??", "??", "??", "??", "??", "??", "??", "??", "??", "??", "??", "??", "??"}
    freq = Counter(g for g in grams if g not in stop)
    return [g for g, c in freq.most_common(top)]


def build_direction_profile(professional, direction, materials=None, profiles=None):
    profiles = profiles or load_profiles()
    directions = profiles.get("directions", {})
    if direction in directions:
        return directions[direction]
    texts = [t for t in (materials or []) if t]
    joined = "".join(texts)
    return {
        "direction": direction,
        "professional": professional,
        "knowledge_points": [],
        "keywords": extract_keywords(joined, top=12),
        "expression_dimensions": [],
    }


def build_student_project_profile(info, materials):
    return {
        "student": info.get("student") or info.get("student_name", ""),
        "professional": info.get("professional", ""),
        "direction": info.get("direction", ""),
        "topic": info.get("topic", ""),
        "station": info.get("station", ""),
        "facts": extract_keywords("".join(materials), top=20),
        "source": materials,
        "features": [],
    }



FOCUS_MARKERS = ["\u65b9\u6848", "\u68c0\u4fee", "\u7ef4\u62a4", "\u6545\u969c", "\u8bbe\u8ba1"]
COMMON_MARKERS = ["\u68c0\u4fee", "\u6545\u969c", "\u65b9\u6848"]


def extract_student_anchors(texts, topic, station, direction_profile=None):
    joined = "".join(texts or [])
    anchors = {
        "topic": topic or "",
        "station": station or "",
        "object": "",
        "devices": [],
        "chapter_focus": [],
        "task_keywords": [],
    }
    if direction_profile:
        devs = [k for k in direction_profile.get("keywords", []) if k and k in joined]
        anchors["devices"] = devs[:6]
    if anchors["devices"]:
        anchors["object"] = (station or "") + anchors["devices"][0]
    words = extract_keywords(joined, top=12)
    anchors["task_keywords"] = words[:6]
    focus = []
    for s in split_sentences(joined):
        if any(m in s for m in FOCUS_MARKERS) and s not in focus:
            focus.append(s)
        if len(focus) >= 3:
            break
    anchors["chapter_focus"] = focus
    return anchors


def rewrite_answer_structured(answer, ctx, profile):
    sentences = split_sentences(answer)
    if len(sentences) < 3:
        return {"text": answer, "changed": False, "reason": "sentence too few"}
    anchors = extract_student_anchors([answer], ctx.get("topic", ""), ctx.get("station", ""), profile)
    lead = None
    if anchors["station"] and anchors["devices"] and anchors["topic"]:
        lead = "%s%s\u662f\u672c\u6b21\u8bbe\u8ba1\u7684\u4e3b\u8981\u5bf9\u8c61\uff0c\u8bfe\u9898\u4e3a\u300a%s\u300b\u3002" % (
            anchors["station"], anchors["devices"][0], anchors["topic"])
    if lead is None or normalize(lead) in [normalize(s) for s in sentences]:
        unique = []
        for s in sentences:
            if ctx.get("station") and ctx.get("station") in s:
                if s not in unique:
                    unique.append(s)
            elif anchors["devices"] and any(d in s for d in anchors["devices"]):
                if s not in unique:
                    unique.append(s)
        lead = unique[0] if unique else sentences[0]
    rest = [s for s in sentences if normalize(s) != normalize(lead)]
    common = []
    design = []
    for s in rest:
        if any(k in s for k in COMMON_MARKERS):
            design.append(s)
        else:
            common.append(s)
    ordered = [lead] + common + design
    new_text = "\n".join(ordered)
    if normalize(new_text) == normalize(answer):
        return {"text": answer, "changed": False, "reason": "no change"}
    return {
        "text": new_text,
        "changed": True,
        "reason": "student anchor lead + common + design structure",
        "anchors": anchors,
    }



ANSWER_PREFIXES = ("\u7b54\uff1a", "\u7b54:")


def extract_defense_answers(docx_path):
    doc = Document(docx_path)
    cell = doc.tables[0].rows[0].cells[0]
    out = []
    for i, p in enumerate(cell.paragraphs):
        t = p.text.strip()
        for pre in ANSWER_PREFIXES:
            if t.startswith(pre):
                out.append({"idx": i, "prefix": pre, "content": t[len(pre):].strip()})
                break
    return out


def _set_paragraph_with_breaks(p, text):
    parts = text.split("\n")
    rpr_src = None
    for r in p.runs:
        rpr_src = r._r.find(qn("w:rPr"))
        if rpr_src is not None:
            break
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    for h in list(p._p.findall(qn("w:hyperlink"))):
        p._p.remove(h)
    for i, part in enumerate(parts):
        if i > 0:
            br_run = OxmlElement("w:r")
            if rpr_src is not None:
                br_run.append(copy.deepcopy(rpr_src))
            br_run.append(OxmlElement("w:br"))
            p._p.append(br_run)
        r = OxmlElement("w:r")
        if rpr_src is not None:
            r.append(copy.deepcopy(rpr_src))
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = part
        r.append(t)
        p._p.append(r)


def apply_defense_answers(docx_path, answers):
    doc = Document(docx_path)
    cell = doc.tables[0].rows[0].cells[0]
    for a in answers:
        if a["idx"] >= len(cell.paragraphs):
            continue
        p = cell.paragraphs[a["idx"]]
        _set_paragraph_with_breaks(p, a["prefix"] + a["content"])
    doc.save(docx_path)



def rewrite_answer_light(answer, ctx, profile):
    sentences = split_sentences(answer)
    if not sentences:
        return {"text": answer, "changed": False, "reason": "empty"}
    anchors = extract_student_anchors([answer], ctx.get("topic", ""), ctx.get("station", ""), profile)
    lead = None
    if anchors["station"] and anchors["devices"] and anchors["topic"]:
        lead = "%s%s\u662f\u672c\u8bbe\u8ba1\u7684\u4e3b\u8981\u5bf9\u8c61\uff0c\u8bfe\u9898\u4e3a\u300a%s\u300b\u3002" % (
            anchors["station"], anchors["devices"][0], anchors["topic"])
    if lead and normalize(lead) not in [normalize(s) for s in sentences]:
        new_text = lead + "\n" + "\n".join(sentences)
        return {"text": new_text, "changed": True, "reason": "light student opener, original kept", "anchors": anchors}
    return {"text": answer, "changed": False, "reason": "no unique anchor"}
