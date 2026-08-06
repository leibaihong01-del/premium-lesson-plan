# -*- coding: utf-8 -*-
"""Defense record template-skeleton normalizer.

The Yang Zhenhai template's 27-paragraph cell is the only layout source:
1. clone the template cell skeleton (paragraph order + paragraph properties)
2. map student Q&A content into question/answer slots
3. preserve fixed regions (header rows, checklist, conclusion, signature)
4. never add or delete template paragraphs
"""
import copy
import json
import os
import re

from docx import Document
from docx.oxml.ns import qn


ENABLED = os.environ.get("DEFENSE_LAYOUT_NORMALIZE", "0") == "1"


def _slot(segment_id, name, slot_type="content"):
    return {"id": segment_id, "type": "slot", "name": name, "slot_type": slot_type}


def build_template_skeleton(template_path):
    """Return segment model of the defense template cell."""
    doc = Document(template_path)
    paras = doc.tables[0].rows[0].cells[0].paragraphs
    segments = []
    for idx, p in enumerate(paras):
        text = p.text.strip()
        seg_id = "P%02d" % idx
        if idx == 2:
            segments.append(_slot(seg_id, "student_info"))
        elif idx == 8:
            segments.append(_slot(seg_id, "question_slot_1"))
        elif idx == 9:
            segments.append(_slot(seg_id, "answer_slot_1"))
        elif idx == 11:
            segments.append(_slot(seg_id, "question_slot_2"))
        elif idx == 12:
            segments.append(_slot(seg_id, "answer_slot_2"))
        elif idx == 23 and text.startswith("\u7B54\u8FA9\u7ED3\u8BBA"):
            segments.append(_slot(seg_id, "conclusion_slot"))
        elif text:
            segments.append({"id": seg_id, "type": "fixed", "content": text})
        else:
            segments.append({"id": seg_id, "type": "spacer"})
    return {"template": "defense_record", "segment_count": len(segments), "segments": segments}


def save_template_skeleton(template_path, output_path):
    skeleton = build_template_skeleton(template_path)
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(skeleton, f, ensure_ascii=False, indent=2)
    return skeleton


def _cell_paragraphs(doc, table_index=0, cell_index=0):
    return doc.tables[table_index].rows[0].cells[cell_index].paragraphs


def _student_info_from_evaluation_table(student_doc):
    """Parse the student info block from the 16x13 evaluation table."""
    if not student_doc.tables or len(student_doc.tables[0].rows) < 8:
        return {"student_name": "", "student_id": "", "class_name": ""}
    table = student_doc.tables[0]
    name = student_id = class_name = ""
    seen = set()
    for row in table.rows:
        cells = []
        for c in row.cells:
            if c._tc in seen:
                continue
            seen.add(c._tc)
            cells.append(c.text.strip())
        for i, text in enumerate(cells):
            if text == "\u59D3\u540D" and i + 1 < len(cells):
                name = cells[i + 1]
            elif text == "\u5B66\u53F7" and i + 1 < len(cells):
                student_id = cells[i + 1]
            elif text == "\u73ED\u7EA7\u540D\u79F0" and i + 1 < len(cells):
                class_name = cells[i + 1]
    return {"student_name": name, "student_id": student_id, "class_name": class_name}


def _student_content_block(student_doc):
    """Extract the Q&A content block from the student defense record."""
    info = _student_info_from_evaluation_table(student_doc)
    target = None
    for table in student_doc.tables:
        if len(table.rows) == 1 and len(table.columns) == 1:
            target = table.rows[0].cells[0]
            break
    if target is None and student_doc.tables:
        target = student_doc.tables[0].rows[0].cells[0]
    if target is None:
        return info
    paras = target.paragraphs
    q1 = q2 = a1 = a2 = None
    conclusion = None
    name = student_id = class_name = ""
    answer_prefix = re.compile(r"^\s*\u7B54\s*[:\uFF1A]")
    for p in paras:
        text = p.text.strip()
        if not text:
            continue
        low = text.lower()
        if text.startswith("\u5B66\u751F\u59D3\u540D"):
            m = re.search(r"\u5B66\u53F7\uFF1A([0-9A-Za-z]+)", text)
            if m:
                student_id = m.group(1)
            name = text.replace("\u5B66\u751F\u59D3\u540D\uFF1A", "").replace("\uFF08\u5B66\u53F7\uFF1A", " ").split(" ")[0].strip()
            cm = re.search(r"\u73ED\u7EA7\uFF1A([^\uFF09)\s]+)", text)
            if cm:
                class_name = cm.group(1)
            continue
        if text.startswith("\u7B54\u8FA9\u7ED3\u8BBA"):
            conclusion = text
            continue
        if answer_prefix.match(text):
            if a1 is None:
                a1 = text
            elif a2 is None:
                a2 = text
            continue
        if re.match(r"^\s*[1-9][0-9]*(\.|\u3001)", text) or "\uFF1F" in text or "?" in text:
            if q1 is None:
                q1 = text
            elif q2 is None:
                q2 = text
            continue
    if not name:
        name = info.get("student_name") or ""
    if not student_id:
        student_id = info.get("student_id") or ""
    if not class_name:
        class_name = info.get("class_name") or ""
    return {
        "student_name": name,
        "student_id": student_id,
        "class_name": class_name,
        "question_1": q1,
        "answer_1": a1,
        "question_2": q2,
        "conclusion": conclusion,
        "answer_2": a2,
    }


def _clear_cell(cell):
    for child in list(cell._tc):
        if child.tag != qn("w:tcPr"):
            cell._tc.remove(child)


SLOT_RUN_INDEX = {
    "student_info": 2,
    "question_slot_1": 0,
    "answer_slot_1": 0,
    "question_slot_2": 0,
    "answer_slot_2": 0,
    "conclusion_slot": 2,
}

SLOT_STRIP_PREFIX = {
    "student_info": "\u5b66\u751f\u59d3\u540d\uff1a",
    "conclusion_slot": "\u7b54\u8fa9\u7ed3\u8bba\uff1a",
}


def _fill_paragraph(target_p, text, run_index=0, strip_prefix=""):
    """Replace the value run text while keeping template run formatting."""
    if not text:
        return
    if strip_prefix and text.startswith(strip_prefix):
        text = text[len(strip_prefix):]
    from docx.text.run import Run
    runs = [Run(r, target_p) for r in target_p._p.r_lst]
    if runs:
        idx = min(run_index, len(runs) - 1)
        runs[idx].text = text
        for run in runs[idx + 1:]:
            run.text = ""
    else:
        target_p.add_run(text)


def normalize(docx_path, template_path, student_path):
    """Rebuild the defense cell using the template skeleton."""
    if not ENABLED or not all(os.path.isfile(p) for p in (docx_path, template_path, student_path)):
        return False
    tpl_doc = Document(template_path)
    student_doc = Document(student_path)
    content = _student_content_block(student_doc)

    out_doc = Document(docx_path)
    out_cell = out_doc.tables[0].rows[0].cells[0]
    skeleton_cell = tpl_doc.tables[0].rows[0].cells[0]
    skeleton_paras = skeleton_cell.paragraphs
    skeleton = build_template_skeleton(template_path)
    segments = skeleton["segments"]

    _clear_cell(out_cell)
    for seg in segments:
        idx = int(seg["id"][1:])
        new_p = copy.deepcopy(skeleton_paras[idx]._p)
        out_cell._tc.append(new_p)

    from docx.text.paragraph import Paragraph
    out_paras = [Paragraph(p._p, out_cell) for p in out_cell.paragraphs]
    slot_values = {
        "student_info": "\u5B66\u751F\u59D3\u540D\uFF1A%s\uFF08\u5B66\u53F7\uFF1A%s\u3001\u73ED\u7EA7\uFF1A%s\uFF09" % (
            content.get("student_name") or "\u5B66\u751F",
            content.get("student_id") or "",
            content.get("class_name") or "",
        ),
        "question_slot_1": content.get("question_1") or "",
        "answer_slot_1": content.get("answer_1") or "",
        "question_slot_2": content.get("question_2") or "",
        "answer_slot_2": content.get("answer_2") or "",
        "conclusion_slot": content.get("conclusion") or "",
    }
    for seg, para in zip(segments, out_paras):
        if seg["type"] == "slot":
            name = seg["name"]
            _fill_paragraph(para, slot_values.get(name, ""), SLOT_RUN_INDEX.get(name, 0), SLOT_STRIP_PREFIX.get(name, ""))

    # Row height follows the template skeleton rule.
    tr_pr = out_doc.tables[0].rows[0]._tr.get_or_add_trPr()
    tr_height = tr_pr.find(qn("w:trHeight"))
    if tr_height is None:
        from docx.oxml import OxmlElement
        tr_height = OxmlElement("w:trHeight")
        tr_pr.append(tr_height)
    tr_height.set(qn("w:val"), "12547")
    tr_height.set(qn("w:hRule"), "atLeast")

    out_doc.save(docx_path)
    return True


def main():
    import sys
    if len(sys.argv) == 4:
        ok = normalize(sys.argv[1], sys.argv[2], sys.argv[3])
        print("normalized" if ok else "skipped")
        return 0
    if len(sys.argv) == 3:
        skeleton = save_template_skeleton(sys.argv[1], sys.argv[2])
        print("skeleton saved:", skeleton["segment_count"])
        return 0
    print("usage: defense_template_normalizer.py <docx> <template> <student_docx>")
    print("       defense_template_normalizer.py <template> <skeleton_json>")
    return 1


if __name__ == "__main__":
    raise SystemExit(main())
