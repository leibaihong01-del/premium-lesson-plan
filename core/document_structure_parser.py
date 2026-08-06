# -*- coding: utf-8 -*-
"""DocumentStructureParser：统一 DOCX/PDF 结构化理解解析器。"""
import hashlib
import os

from docx import Document
from docx.oxml.ns import qn

from document_field_extractor import DocumentFieldExtractor
from pdf_layout_evidence import extract_pdf_evidence
from document_structure import empty_document_structure


def _p_text(el):
    return "".join(t.text or "" for t in el.xpath(".//w:t")).strip()


def _parse_docx(path):
    doc = Document(path)
    paragraphs = []
    styles = {}
    full_parts = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        style = p.style.name if p.style else ""
        styles[style] = styles.get(style, 0) + 1
        pf = p.paragraph_format
        runs = []
        for r in p.runs:
            if not r.text.strip():
                continue
            rPr = r._r.find(qn("w:rPr"))
            east_asia = None
            if rPr is not None:
                rFonts = rPr.find(qn("w:rFonts"))
                if rFonts is not None:
                    east_asia = rFonts.get(qn("w:eastAsia"))
            runs.append({
                "text": r.text,
                "bold": bool(r.font.bold),
                "size_pt": r.font.size.pt if r.font.size else None,
                "font": r.font.name,
                "east_asia": east_asia,
            })
        paragraphs.append({
            "index": i,
            "type": "paragraph",
            "text": text,
            "page": None,
            "style": style,
            "alignment": str(p.alignment) if p.alignment is not None else None,
            "line_spacing": pf.line_spacing,
            "left_indent_emu": pf.left_indent,
            "first_line_indent_emu": pf.first_line_indent,
            "runs": runs,
        })
        if text:
            full_parts.append(text)

    tables = []
    for ti, t in enumerate(doc.tables):
        grid = [gc.get(qn("w:w")) for gc in t._tbl.findall(qn("w:tblGrid"))]
        merge_sig = []
        row_cells = []
        for ri, row in enumerate(t.rows):
            seen = set()
            cells = []
            sig = ""
            for ci, c in enumerate(row.cells):
                if c._tc in seen:
                    sig += "M"
                    continue
                seen.add(c._tc)
                sig += "U"
                text = c.text.strip()
                cells.append({
                    "row": ri,
                    "col": ci,
                    "text": text,
                    "paragraphs": [pp.text for pp in c.paragraphs],
                })
                if text:
                    full_parts.append(text)
            merge_sig.append(sig)
            row_cells.append(cells)
        tables.append({
            "table_id": ti,
            "type": "table",
            "rows_count": len(t.rows),
            "cols_count": len(t.columns),
            "grid": grid,
            "merge_sig": merge_sig,
            "row_cells": row_cells,
        })

    return {
        "paragraphs": paragraphs,
        "tables": tables,
        "styles": styles,
        "full_parts": full_parts,
        "sections_count": len(doc.sections),
        "toc_field_present": bool(doc.element.body.findall(".//" + qn("w:fldSimple"))) or bool(
            doc.element.body.findall(".//" + qn("w:instrText"))),
    }


def _parse_pdf(path):
    if not path or not os.path.isfile(path):
        return {"count": None, "per_page_chars": [], "error": "pdf_missing"}
    result = {"count": None, "per_page_chars": []}
    try:
        from pypdf import PdfReader
        result["count"] = len(PdfReader(path).pages)
    except Exception:
        result["count"] = None
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                result["per_page_chars"].append(len(page.chars))
    except Exception as e:
        result["pdfplumber_error"] = repr(e)
    return result


def parse(path, pdf_path=None, document_type=""):
    structure = empty_document_structure(document_type=document_type, source_path=path, pdf_path=pdf_path)
    raw = _parse_docx(path)
    structure.update({
        "paragraphs": raw["paragraphs"],
        "tables": raw["tables"],
        "styles": {"counts": raw["styles"]},
        "sections_count": raw["sections_count"],
        "toc_field_present": raw["toc_field_present"],
        "full_text": "\n".join(raw["full_parts"]),
        "pages": _parse_pdf(pdf_path),
    })
    structure["document_id"] = hashlib.sha256(open(path, "rb").read()).hexdigest()[:12]
    DocumentFieldExtractor().extract(structure)
    return structure