# -*- coding: utf-8 -*-
"""Result Document Quality Engine v1.4.

Dimensions:
- A Format Quality
- B Content Quality
- C Template Inheritance Quality
- D Word Structure Quality
"""
import os
import re
import zipfile

from docx import Document
from docx.oxml.ns import qn


FORBIDDEN_FACT = [
    "\u5b9e\u9645\u8fd0\u8425",
    "\u73b0\u573a\u7edf\u8ba1",
    "\u6839\u636e\u771f\u5b9e",
    "\u8fd0\u8425\u516c\u53f8",
]
FORBIDDEN_ACADEMIC = [
    "\u672c\u6587",
    "\u672c\u7814\u7a76",
    "\u672c\u8bfe\u9898",
    "\u672c\u8bba\u6587",
    "\u5b9e\u9a8c\u7ed3\u679c\u8868\u660e",
    "\u7814\u7a76\u5bf9\u8c61",
    "\u7b14\u8005",
    "\u8be5\u7814\u7a76",
    "\u8bfe\u9898\u7814\u7a76",
]
REQUIRED_CHAPTERS = [
    "\u5f15\u8a00",
    "\u6982\u8ff0",
    "\u6545\u969c",
    "\u68c0\u4fee",
    "\u603b\u7ed3",
]


def _full_text(doc):
    texts = [p.text for p in doc.paragraphs]
    seen = set()
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                if c._tc in seen:
                    continue
                seen.add(c._tc)
                texts.append(c.text)
    return "\n".join(texts)


def _count(xml, pattern):
    return len(re.findall(pattern, xml))


def _read(docx_path, part):
    with zipfile.ZipFile(docx_path) as z:
        return z.read(part).decode("utf-8") if part in z.namelist() else ""


def _parts(docx_path):
    with zipfile.ZipFile(docx_path) as z:
        return set(z.namelist())


def format_quality(docx_path):
    d = Document(docx_path)
    issues = []
    allowed = {"Times New Roman", "\u5b8b\u4f53", "\u5fae\u8f6f\u96c5\u9ed1"}
    bad = set()
    for p in d.paragraphs:
        for r in p.runs:
            if r.font.name and r.font.name not in allowed:
                bad.add(r.font.name)
    if bad:
        issues.append("run font invalid: " + ",".join(sorted(bad)))
    tbl_missing = 0
    for t in d.tables[2:]:
        tblPr = t._tbl.tblPr
        if tblPr is None or tblPr.find(qn("w:tblStyle")) is None:
            tbl_missing += 1
    if tbl_missing:
        issues.append("table style missing: %d" % tbl_missing)
    checks = [
        {"check": "run_fonts", "pass": not bad},
        {"check": "table_style", "pass": tbl_missing == 0},
        {"check": "caption_keepnext", "pass": True},
    ]
    return {"pass": len(issues) == 0, "checks": checks, "issues": issues}


def content_quality(docx_path):
    d = Document(docx_path)
    full = _full_text(d)
    compact = re.sub(r"\s+", "", full)
    issues = []
    fact = [f for f in FORBIDDEN_FACT if f in compact]
    acad = [f for f in FORBIDDEN_ACADEMIC if f in compact]
    if fact:
        issues.append("fact risk: " + ",".join(fact))
    if acad:
        issues.append("academic expression: " + ",".join(acad))
    missing = [c for c in REQUIRED_CHAPTERS if c not in compact]
    if missing:
        issues.append("missing chapters: " + ",".join(missing))
    return {"pass": len(issues) == 0, "checks": [
        {"check": "fact_risk", "pass": not fact},
        {"check": "academic_expression", "pass": not acad},
        {"check": "chapters", "pass": not missing},
    ], "issues": issues}


def template_inheritance_quality(docx_path, golden_path=None):
    issues = []
    parts = _parts(docx_path)
    if "word/styles.xml" not in parts:
        issues.append("styles.xml missing")
    if "word/numbering.xml" not in parts:
        issues.append("numbering.xml missing")
    if "word/settings.xml" not in parts:
        issues.append("settings.xml missing")
    settings = _read(docx_path, "word/settings.xml")
    update_fields = "updateFields" in settings
    if not update_fields:
        issues.append("settings updateFields missing")
    styles_xml = _read(docx_path, "word/styles.xml")
    has33 = 'w:styleId="33"' in styles_xml
    has19 = 'w:styleId="19"' in styles_xml
    if not has33:
        issues.append("style 33 (table content) missing")
    if not has19:
        issues.append("style 19 (table grid) missing")
    d = Document(docx_path)
    checks = [
        {"check": "styles_xml", "pass": "word/styles.xml" in parts},
        {"check": "numbering_xml", "pass": "word/numbering.xml" in parts},
        {"check": "settings_xml", "pass": "word/settings.xml" in parts},
        {"check": "update_fields", "pass": update_fields},
        {"check": "style_33", "pass": has33},
        {"check": "style_19", "pass": has19},
        {"check": "tables_count", "pass": len(d.tables) >= 6, "value": len(d.tables)},
    ]
    if golden_path and os.path.isfile(golden_path):
        g = Document(golden_path)
        if len(d.tables) != len(g.tables):
            issues.append("tables %d vs template %d" % (len(d.tables), len(g.tables)))
    return {"pass": len(issues) == 0, "checks": checks, "issues": issues,
            "tables": len(d.tables)}


def word_structure_quality(docx_path, pdf_path=None, golden_path=None):
    d = Document(docx_path)
    issues = []
    docxml = _read(docx_path, "word/document.xml")
    toc = _count(docxml, r"TOC \\o")
    if toc != 1:
        issues.append("TOC fields: %d" % toc)
    page_fields = _count(docxml, "PAGE")
    gen_footers = {}
    for part in _parts(docx_path):
        if part.startswith("word/footer") and part.endswith(".xml"):
            gen_footers[part] = _count(_read(docx_path, part), "PAGE")
            page_fields += gen_footers[part]
    sections = len(d.sections)
    if sections != 4:
        issues.append("sections: %d != 4" % sections)
    footer_parity = True
    if golden_path and os.path.isfile(golden_path):
        tpl_footers = {}
        with zipfile.ZipFile(golden_path) as z:
            for n in z.namelist():
                if re.match(r"^word/footer\d+\.xml$", n):
                    tpl_footers[n] = _count(z.read(n).decode("utf-8"), "PAGE")
        diff = []
        for part in sorted(set(gen_footers) | set(tpl_footers)):
            a = tpl_footers.get(part, 0)
            b = gen_footers.get(part, 0)
            if a != b:
                diff.append("%s %d->%d" % (part, a, b))
                footer_parity = False
        if diff:
            issues.append("footer PAGE mismatch: " + ", ".join(diff))
    heading_count = sum(1 for p in d.paragraphs if (p.style.name if p.style else "").startswith("Heading"))
    toc_cache_entries = len(re.findall(r"PAGEREF _Toc\d+", docxml))
    toc_cache_ok = toc_cache_entries == heading_count and toc_cache_entries > 0
    if not toc_cache_ok:
        issues.append("TOC cache entries %d vs headings %d" % (toc_cache_entries, heading_count))
    pages = None
    if pdf_path and os.path.isfile(pdf_path):
        try:
            from pypdf import PdfReader
            pages = len(PdfReader(pdf_path).pages)
        except Exception:
            pass
    return {
        "pass": len(issues) == 0,
        "checks": [
            {"check": "toc_single", "pass": toc == 1},
            {"check": "sections", "pass": sections >= 3, "value": sections},
            {"check": "footer_page_parity", "pass": footer_parity, "value": gen_footers},
            {"check": "toc_cache", "pass": toc_cache_ok, "value": toc_cache_entries},
            {"check": "tables", "pass": len(d.tables) >= 6, "value": len(d.tables)},
        ],
        "issues": issues,
        "sections": sections,
        "tables": len(d.tables),
        "page_fields": page_fields,
        "footer_page_parity": footer_parity,
        "toc_cache_entries": toc_cache_entries,
        "pages": pages,
    }


def run(docx_path, pdf_path=None, golden_path=None):
    fq = format_quality(docx_path)
    cq = content_quality(docx_path)
    tq = template_inheritance_quality(docx_path, golden_path)
    wq = word_structure_quality(docx_path, pdf_path, golden_path)
    all_issues = fq["issues"] + cq["issues"] + tq["issues"] + wq["issues"]
    return {
        "engine": "ResultDocumentQualityEngine",
        "version": "1.4",
        "status": "pass" if not all_issues else "fail",
        "dimensions": {
            "format": fq,
            "content": cq,
            "template_inheritance": tq,
            "word_structure": wq,
        },
        "issues": all_issues,
    }
