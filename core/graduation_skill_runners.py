# -*- coding: utf-8 -*-
"""V0.7 Graduation Skill Runners：四类文档统一走 StudentProfile → Experience → DocumentStructure → Quality Sense。"""
import copy
import json
import os
import re
import shutil
import sys
import time

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.oxml.ns import qn


def _unique_cells(row):
    seen = set()
    out = []
    for c in row.cells:
        if c._tc in seen:
            continue
        seen.add(c._tc)
        out.append(c)
    return out


def _set_para(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def _replace_runs(doc, pairs):
    def fix(paras):
        for p in paras:
            for r in p.runs:
                for old, new in pairs:
                    if old in r.text:
                        r.text = r.text.replace(old, new)
    fix(doc.paragraphs)
    seen = set()
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                if c._tc in seen:
                    continue
                seen.add(c._tc)
                fix(c.paragraphs)


def _render_pdf(docx_path, pdf_path, v06_dir):
    sys.path.insert(0, v06_dir)
    from render_docx import render_to_pdf
    return render_to_pdf(docx_path, pdf_path) and os.path.isfile(pdf_path)


def _write_json(path, payload):
    os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=2)


def _station(topic):
    m = re.search(r"([\u4e00-\u9fa5]{2,8}?站)", topic or "")
    return m.group(1) if m else "黄兴南路站"


def _taskbook_experience_trace(student, process_dir, courseagent_root, project_root):
    sys.path.insert(0, os.path.join(courseagent_root, "core"))
    from experience_loader import ExperienceLoader
    loader = ExperienceLoader(enabled=True, courseagent_root=courseagent_root)
    applicable = loader.load("task_book", template="01 杨振海 毕业设计任务书 ...docx", task_context={"student": student})
    payload = {
        "trace_version": "0.7-trace-v1",
        "document_type": "task_book",
        "skill": "TaskBookSkillRunner",
        "student": student,
        "experience_loaded": [e["experience_id"] for e in applicable["experiences"] if e.get("loaded")],
        "experience_missing": [e["experience_id"] for e in applicable["experiences"] if not e.get("loaded")],
    }
    _write_json(os.path.join(process_dir, "experience_trace_task_book.json"), payload)
    return payload


def _result_experience_trace(student, process_dir, courseagent_root, project_root):
    sys.path.insert(0, os.path.join(courseagent_root, "core"))
    from result_experience_consumer import ResultExperienceConsumer
    consumer = ResultExperienceConsumer(enabled=True, courseagent_root=courseagent_root)
    ctx = consumer.build_context(student, "02 杨振海 毕业设计成果 ...docx")
    payload = {
        "trace_version": "0.7-result-experience-trace-v1",
        "document_type": "result",
        "skill": "ResultSkillRunner",
        "student": student,
        "experience_loaded": [e["name"] for e in ctx.loaded_experience],
        "experience_missing": [m["name"] for m in ctx.missing_experience],
    }
    _write_json(os.path.join(process_dir, "experience_trace_result.json"), payload)
    return payload


def _admin_experience_trace(doc_type, skill, student, process_dir, courseagent_root):
    sys.path.insert(0, os.path.join(courseagent_root, "core"))
    from experience_loader import ExperienceLoader
    loader = ExperienceLoader(enabled=True, courseagent_root=courseagent_root)
    applicable = loader.load(doc_type, template="template://%s" % doc_type, task_context={"student": student})
    payload = {
        "trace_version": "0.7-trace-v1",
        "document_type": doc_type,
        "skill": skill,
        "student": student,
        "experience_loaded": [e["experience_id"] for e in applicable["experiences"] if e.get("loaded")],
        "experience_missing": [e["experience_id"] for e in applicable["experiences"] if not e.get("loaded")],
    }
    _write_json(os.path.join(process_dir, "experience_trace_%s.json" % doc_type), payload)
    return payload


def _generation_trace(doc_type, skill, student, template, process_dir, quality_checks=None, revision_actions=None):
    payload = {
        "trace_version": "0.7-generation-trace-v1",
        "document_type": doc_type,
        "skill": skill,
        "student": student,
        "template_source": template,
        "generated_at": time.strftime("%Y-%m-%dT%H:%M:%S"),
        "experience_integration_enabled": True,
        "quality_checks": quality_checks or [],
        "revision_actions": revision_actions or [],
        "final_validation": "generated",
    }
    _write_json(os.path.join(process_dir, "generation_trace_%s.json" % doc_type), payload)
    return payload


class TaskBookSkillRunner:
    name = "task_book_skill"

    def __init__(self, profile, ws, courseagent_root, project_root):
        self.profile = profile
        self.ws = ws
        self.v06 = os.path.join(ws, "00_系统配置", "模块", "v06")
        self.courseagent_root = courseagent_root
        self.project_root = project_root

    def run(self, student_dir, package_dir, process_dir):
        info = self.profile.to_dict()
        topic = self.profile.topic
        template = os.path.join(self.ws, "02_模板文件", "01 杨振海 毕业设计任务书 黄兴南路站AFC闸机设备检修方案设计.docx")
        output = os.path.join(package_dir, "01 %s 毕业设计任务书 %s.docx" % (self.profile.student_name, topic))
        pdf = os.path.join(process_dir, os.path.basename(output).replace(".docx", ".pdf"))
        shutil.copy2(template, output)
        doc = Document(output)
        t = doc.tables[0]
        cells = _unique_cells(t.rows[0])
        _set_para(cells[3].paragraphs[0], self.profile.student_name)
        if len(cells) > 5:
            _set_para(cells[5].paragraphs[0], self.profile.advisor)
        c1 = _unique_cells(t.rows[1])
        _set_para(c1[1].paragraphs[0], self.profile.class_name)
        _set_para(c1[3].paragraphs[0], self.profile.student_id)
        c2 = _unique_cells(t.rows[2])
        _set_para(c2[1].paragraphs[0], topic)

        src = Document(os.path.join(student_dir, "任务书.docx"))
        st = src.tables[0]
        for ri in (4, 5):
            sc = _unique_cells(st.rows[ri])[1]
            dc = _unique_cells(t.rows[ri])[1]
            for i, p in enumerate(sc.paragraphs):
                if i < len(dc.paragraphs):
                    _set_para(dc.paragraphs[i], p.text)
                else:
                    newp = copy.deepcopy(dc.paragraphs[-1]._p)
                    dc.paragraphs[-1]._p.addnext(newp)
                    _set_para(dc.paragraphs[i], p.text)
        doc.save(output)
        _render_pdf(output, pdf, self.v06)

        markers = {"p1": [], "p2": []}
        try:
            import pdfplumber
            with pdfplumber.open(pdf) as p:
                for i, page in enumerate(p.pages):
                    for key in ["主要设计目标", "主要完成以下任务", "预期成果", "设计进程", "签名"]:
                        if page.search(key):
                            markers["p%d" % (i + 1)].append(key)
        except Exception:
            pass
        page_semantic = {
            "page1_ok": "主要设计目标" in markers["p1"] and "主要完成以下任务" in markers["p1"],
            "page2_ok": all(k in markers["p2"] for k in ["预期成果", "设计进程", "签名"]),
        }
        _taskbook_experience_trace(self.profile.student_name, process_dir, self.courseagent_root, self.project_root)
        _generation_trace("task_book", self.name, self.profile.student_name, os.path.basename(template),
                          process_dir, quality_checks=[{"check": "page_semantic", "result": page_semantic}])
        return {"skill": self.name, "output": output, "pdf": pdf, "page_semantic": page_semantic}


def _transplant_cover(dst_doc, src_doc):
    import re as _re
    def norm(s):
        return _re.sub(r"[\s:：（）()【】\[\]]", "", s or "")
    def unique_cells(row):
        seen = set()
        out = []
        for c in row.cells:
            if c._tc in seen:
                continue
            seen.add(c._tc)
            out.append(c)
        return out
    def find(table, label):
        key = norm(label)
        for row in table.rows:
            cells = unique_cells(row)
            for i, c in enumerate(cells[:-1]):
                if key in norm(c.text) and c.text.strip():
                    return cells[i + 1]
        return None
    if not dst_doc.tables or not src_doc.tables:
        return
    for label in ["选题名称", "选题类型", "二级学院", "专业名称", "班级名称", "姓名", "学号", "指导教师", "完成时间"]:
        sv = find(src_doc.tables[0], label)
        dv = find(dst_doc.tables[0], label)
        if sv is None or dv is None:
            continue
        tc = dv._tc
        for child in list(tc):
            if child.tag != qn("w:tcPr"):
                tc.remove(child)
        for p in sv.paragraphs:
            tc.append(copy.deepcopy(p._p))


def _normalize_result_fonts(doc):
    for p in doc.paragraphs:
        style = p.style.name if p.style else ""
        if not p.text.strip():
            continue
        if "Heading 1" in style:
            size, bold, east = 16, True, "黑体"
        elif "Heading 2" in style:
            size, bold, east = 15, True, "黑体"
        else:
            size, bold, east = 12, False, "宋体"
        for r in p.runs:
            r.font.size = Pt(size)
            r.font.bold = bold
            r.font.name = "Times New Roman"
            rPr = r._r.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.append(rFonts)
            rFonts.set(qn("w:ascii"), "Times New Roman")
            rFonts.set(qn("w:hAnsi"), "Times New Roman")
            rFonts.set(qn("w:eastAsia"), east)




def _apply_reference_format_experience(doc):
    """应用已验证经验：邱志豪参考文献格式修正（仅调整段落格式，不改内容）。"""
    in_ref = False
    ref_paras = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t == "参考文献":
            in_ref = True
            continue
        if in_ref and t:
            ref_paras.append(p)
    for p in ref_paras:
        pPr = p._p.get_or_add_pPr()
        for ind in pPr.findall(qn("w:ind")):
            pPr.remove(ind)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "420")
        ind.set(qn("w:hanging"), "420")
        pPr.append(ind)
        for sp in pPr.findall(qn("w:spacing")):
            pPr.remove(sp)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:line"), "440")
        spacing.set(qn("w:lineRule"), "exact")
        pPr.append(spacing)
    return len(ref_paras)

class ResultSkillRunner:
    name = "result_skill"

    def __init__(self, profile, ws, courseagent_root, project_root):
        self.profile = profile
        self.ws = ws
        self.v03 = os.path.join(ws, "00_系统配置", "模块", "v03")
        self.v06 = os.path.join(ws, "00_系统配置", "模块", "v06")
        self.courseagent_root = courseagent_root
        self.project_root = project_root

    def run(self, student_dir, package_dir, process_dir):
        sys.path.insert(0, self.v03)
        import result_reference_builder as rrb
        info = self.profile.to_dict()
        topic = self.profile.topic
        template = os.path.join(self.ws, "02_模板文件", "02 杨振海 毕业设计成果 黄兴南路站AFC闸机设备检修方案设计.docx")
        output = os.path.join(package_dir, "02 %s 毕业设计成果 %s.docx" % (self.profile.student_name, topic))
        pdf = os.path.join(process_dir, os.path.basename(output).replace(".docx", ".pdf"))
        draft = os.path.join(student_dir, "成果初稿.docx")
        rrb.reconstruct(self.profile.student_name, self.profile.direction, "02", info, draft, output, neutralize=True)
        doc = Document(output)
        cover = os.path.join(student_dir, "成果记录表.docx")
        if os.path.isfile(cover):
            _transplant_cover(doc, Document(cover))
        _normalize_result_fonts(doc)
        ref_count = _apply_reference_format_experience(doc)
        doc.save(output)
        self._write_reference_experience_trace(process_dir, ref_count)
        _render_pdf(output, pdf, self.v06)
        _result_experience_trace(self.profile.student_name, process_dir, self.courseagent_root, self.project_root)
        quality = self._run_quality(package_dir, process_dir, output, pdf, topic)
        _generation_trace("result", self.name, self.profile.student_name, os.path.basename(template),
                          process_dir, quality_checks=[{"check": "result_quality", "status": quality.get("quality_status")}])
        return {"skill": self.name, "output": output, "pdf": pdf, "quality": quality}



    @staticmethod
    def _write_reference_experience_trace(process_dir, ref_count):
        import time as _t
        payload = {
            "trace_version": "0.7-result-reference-experience-trace-v1",
            "experience": "邱志豪参考文献格式修正经验",
            "stage": "Result Generation / Result Quality Refinement",
            "content": "参考文献格式规范化（悬挂缩进 0.74cm/420twips，行距 440 exact）",
            "applied_paragraphs": ref_count,
            "result": "applied",
            "generated_at": _t.strftime("%Y-%m-%dT%H:%M:%S"),
        }
        path = os.path.join(process_dir, "result_reference_experience_trace.json")
        with open(path, "w", encoding="utf-8") as f:
            json.dump(payload, f, ensure_ascii=False, indent=2)
        return path

    def _run_quality(self, package_dir, process_dir, output, pdf, topic):
        sys.path.insert(0, os.path.join(self.courseagent_root, "core"))
        from result_quality_pipeline import ResultQualityPipeline
        from result_experience_consumer import ResultExperienceConsumer
        consumer = ResultExperienceConsumer(enabled=True, courseagent_root=self.courseagent_root)
        ctx = consumer.build_context(self.profile.student_name, "02 杨振海 毕业设计成果 ...docx")
        taskbook = os.path.join(package_dir, "01 %s 毕业设计任务书 %s.docx" % (self.profile.student_name, topic))
        pipeline = ResultQualityPipeline(output_dir=process_dir)
        result = pipeline.run(profile=self.profile, document_path=output, pdf_path=pdf,
                              template_path=os.path.join(self.ws, "02_模板文件",
                                                         "02 杨振海 毕业设计成果 黄兴南路站AFC闸机设备检修方案设计.docx"),
                              taskbook_path=taskbook, experience_context=ctx)
        return {"quality_status": result["quality_status"], "trace": result["quality_pipeline_trace"]}


class EvaluationSkillRunner:
    name = "evaluation_form_skill"

    def __init__(self, profile, ws, courseagent_root, project_root):
        self.profile = profile
        self.ws = ws
        self.v06 = os.path.join(ws, "00_系统配置", "模块", "v06")
        self.courseagent_root = courseagent_root
        self.project_root = project_root

    def run(self, student_dir, package_dir, process_dir):
        topic = self.profile.topic
        template = os.path.join(self.ws, "02_模板文件", "04 杨振海 毕业设计成绩评定表 黄兴南路站AFC闸机设备检修方案设计.docx")
        output = os.path.join(package_dir, "03 %s 毕业设计成绩评定表 %s.docx" % (self.profile.student_name, topic))
        pdf = os.path.join(process_dir, os.path.basename(output).replace(".docx", ".pdf"))
        shutil.copy2(template, output)
        doc = Document(output)
        pairs = [
            ("杨振海", self.profile.student_name),
            ("202421044622", self.profile.student_id),
            ("24级机电技术1班", self.profile.class_name),
            ("黄兴南路站AFC闸机设备检修方案设计", topic),
            ("黄兴南路站", _station(topic)),
        ]
        _replace_runs(doc, pairs)
        doc.save(output)
        _render_pdf(output, pdf, self.v06)
        _admin_experience_trace("evaluation_form", self.name, self.profile.student_name, process_dir, self.courseagent_root)
        _generation_trace("evaluation_form", self.name, self.profile.student_name, os.path.basename(template), process_dir)
        return {"skill": self.name, "output": output, "pdf": pdf}


class DefenseSkillRunner:
    name = "defense_record_skill"

    def __init__(self, profile, ws, courseagent_root, project_root):
        self.profile = profile
        self.ws = ws
        self.v06 = os.path.join(ws, "00_系统配置", "模块", "v06")
        self.courseagent_root = courseagent_root
        self.project_root = project_root

    def run(self, student_dir, package_dir, process_dir):
        topic = self.profile.topic
        template = os.path.join(self.ws, "02_模板文件", "05 杨振海 毕业设计答辩记录表 黄兴南路站AFC闸机设备检修方案设计.docx")
        output = os.path.join(package_dir, "04 %s 毕业设计答辩记录表 %s.docx" % (self.profile.student_name, topic))
        pdf = os.path.join(process_dir, os.path.basename(output).replace(".docx", ".pdf"))
        shutil.copy2(template, output)
        doc = Document(output)
        pairs = [
            ("杨振海", self.profile.student_name),
            ("202421044622", self.profile.student_id),
            ("黄兴南路站AFC闸机设备检修方案设计", topic),
            ("黄兴南路站", _station(topic)),
        ]
        _replace_runs(doc, pairs)
        src = os.path.join(student_dir, "答辩记录表.docx")
        if os.path.isfile(src):
            try:
                sdoc = Document(src)
                scell = sdoc.tables[1].rows[0].cells[0]
                dcell = doc.tables[0].rows[0].cells[0]
                for child in list(dcell._tc):
                    if child.tag != qn("w:tcPr"):
                        dcell._tc.remove(child)
                for p in scell.paragraphs:
                    dcell._tc.append(copy.deepcopy(p._p))
                for p in dcell.paragraphs:
                    t = p.text.strip()
                    if t.startswith("学生姓名"):
                        _set_para(p, "学生姓名：%s（学号：%s 班级：%s）" % (
                            self.profile.student_name, self.profile.student_id, self.profile.class_name))
                    elif t.startswith("指导教师"):
                        _set_para(p, "指导教师：%s" % self.profile.advisor)
            except Exception:
                pass
        doc.save(output)
        core_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        sys.path.insert(0, core_dir)
        try:
            from defense_template_normalizer import normalize as normalize_defense_layout
            normalize_defense_layout(output, template, src)
        except Exception:
            pass
        _admin_experience_trace("defense_record", self.name, self.profile.student_name, process_dir, self.courseagent_root)
        _generation_trace("defense_record", self.name, self.profile.student_name, os.path.basename(template), process_dir)
        return {"skill": self.name, "output": output, "pdf": pdf}
