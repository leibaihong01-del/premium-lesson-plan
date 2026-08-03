# -*- coding: utf-8 -*-
"""毕业设计成果区域化内容审查执行器。

规则来源：school_rules（学校硬约束） + expert_rules（教师经验增强），不维护独立违禁词库。
"""
import json
import os
import re

from docx import Document

BASE = os.path.dirname(os.path.abspath(__file__))
GD = os.path.normpath(os.path.join(BASE, "..", ".."))
SCHOOL = os.path.join(GD, "school_rules", "content_rules.json")
EXPERT = os.path.join(GD, "expert_rules", "content_rules.json")


def _load(path):
    with open(path, encoding="utf-8") as f:
        return json.load(f)


class ResultContentValidator:
    def __init__(self, school_path=None, expert_path=None):
        self.school = _load(school_path or SCHOOL)
        self.expert = _load(expert_path or EXPERT)
        self.region_control = self.school.get("region_control", {})
        self.levels = {"forbidden": [], "warning": []}
        for src in (self.school, self.expert):
            for level in ("forbidden", "warning"):
                for rule in src.get("expression_level", {}).get(level, []):
                    self.levels[level].append({**rule, "source": src.get("source", "")})

    def identify_region(self, doc):
        regions = {}
        current = "cover"
        ref_started = False
        for i, p in enumerate(doc.paragraphs):
            t = p.text.strip()
            if ref_started:
                regions[i] = "references"
                continue
            if current == "cover" and "真实性承诺" in t:
                current = "declaration"
            elif current == "declaration" and ("目" in t and "录" in t):
                current = "toc"
            elif current in ("toc", "cover") and i > 0 and re.match(r"^1\s*\S", t):
                current = "body"
            elif current == "toc" and t.startswith("摘要"):
                current = "abstract"
            elif current == "abstract" and re.match(r"^1\s*\S", t):
                current = "body"
            if current == "body" and t == "参考文献":
                ref_started = True
                regions[i] = "references"
                continue
            if current == "body" and re.match(r"^\d+\s*(总结|结论)", t):
                current = "conclusion"
            if current == "body" and t.startswith("附录"):
                current = "appendix"
            regions[i] = current
        return regions

    def analyze(self, docx_path):
        doc = Document(docx_path)
        regions = self.identify_region(doc)
        issues = []
        for i, p in enumerate(doc.paragraphs):
            region = regions.get(i, "unknown")
            if self.region_control.get(region) != "check":
                continue
            style = p.style.name if p.style else ""
            text = p.text
            forbidden_hits = [r for r in self.levels["forbidden"] if r["expression"] in text]
            for r in forbidden_hits:
                issues.append(self._issue(i, region, style, text, r, "forbidden"))
            if forbidden_hits:
                continue
            for r in self.levels["warning"]:
                if r["expression"] in text:
                    issues.append(self._issue(i, region, style, text, r, "warning"))
        return issues

    @staticmethod
    def _issue(i, region, style, text, rule, level):
        return {
            "position": i + 1,
            "region": region,
            "level": level,
            "is_heading": "Heading" in style,
            "original_expression": rule["expression"],
            "context": text.strip()[:80],
            "suggestion": "人工确认" if "Heading" in style else rule.get("suggestion", ""),
            "reason": rule.get("reason", ""),
            "rule_source": rule.get("source", ""),
            "school_rule_source": "08轨道车辆学院毕业设计检查标准",
            "violation": level == "forbidden"
        }

    def render(self, docx_path, issues):
        lines = ["# 毕业设计成果内容规范检查报告", "", "文件：%s" % docx_path, ""]
        if not issues:
            lines.append("未发现区域化内容规范问题。")
        else:
            lines.append("| 位置 | 区域 | 级别 | 原文 | 上下文 | 是否违规 | 建议 | 规则来源 |")
            lines.append("|---|---|---|---|---|---|---|---|")
            for it in issues:
                lines.append("| %s | %s | %s | %s | %s | %s | %s | %s |" % (
                    it["position"], it["region"], it["level"], it["original_expression"],
                    it["context"], "是" if it["violation"] else "待判断",
                    it["suggestion"], it["rule_source"]))
        return "\n".join(lines)
