# -*- coding: utf-8 -*-
"""毕业设计审查 Agent：结构/模板/内容/专业/格式五维检查。"""
import json
import os
import re

from docx import Document
from docx.oxml.ns import qn

WS = os.environ.get("GRAD_WS", r"D:\Users\leibaihong\Desktop\课程材料优化\毕业设计智能制作工作区")
SAMPLE_DIR = os.path.join(WS, "06_输出成果", "周志贤完整样例")
QUALITY = os.path.join(WS, "05_质量检查")
BLUEPRINT = os.path.join(WS, "04_过程记录", "规划蓝图", "周志贤_方案蓝图.json")
NAME = "周志贤"
TOPIC = "周南中学站中央空调日常维护方案设计"
FILES = ["毕业设计任务书", "毕业设计成果", "毕业设计答辩记录表", "毕业设计成绩评定表"]
BANNED = ["杨振海", "202421044622", "AFC", "闸机", "黄兴南路站"]
PLACEHOLDERS = ["待完善", "待确认", "XXX", "后续添加"]
REQUIRED = ["引言", "概述", "故障", "方案", "总结", "参考文献"]


def stats(path):
    doc = Document(path)
    images = 0
    parts = [doc.element.body] + [s.header._element for s in doc.sections] + [s.footer._element for s in doc.sections]
    for part in parts:
        images += len(list(part.iter(qn("w:drawing")))) + len(list(part.iter(qn("w:pict"))))
    page_breaks = 0
    for p in doc.element.body.iter(qn("w:p")):
        pPr = p.find(qn("w:pPr"))
        if pPr is not None:
            if pPr.find(qn("w:sectPr")) is not None:
                page_breaks += 1
            if pPr.find(qn("w:pageBreakBefore")) is not None:
                page_breaks += 1
    for br in doc.element.body.iter(qn("w:br")):
        if br.get(qn("w:type")) == "page":
            page_breaks += 1
    max_empty = 0
    run = 0
    for p in doc.paragraphs:
        if not p.text.strip():
            run += 1
            max_empty = max(max_empty, run)
        else:
            run = 0
    breaks = sum(1 for _ in doc.element.body.iter(qn("w:br")))
    return {"sections": len(doc.sections), "tables": len(doc.tables), "images": images,
            "page_breaks": page_breaks, "max_empty": max_empty, "breaks": breaks}


def review_file(label):
    path = os.path.join(SAMPLE_DIR, "%s_%s_%s.docx" % (NAME, label, TOPIC))
    doc = Document(path)
    text = "\n".join([p.text for p in doc.paragraphs] +
                     [c.text for t in doc.tables for row in t.rows for c in row.cells])
    errors = []
    warnings = []
    # 内容
    hits = [b for b in BANNED if b in text]
    if hits:
        errors.append({"type": "content", "problem": "模板污染：" + "、".join(hits)})
    ph = [p for p in PLACEHOLDERS if p in text]
    if ph:
        errors.append({"type": "content", "problem": "占位符：" + "、".join(ph)})
    if NAME not in text or "202421044804" not in text:
        errors.append({"type": "content", "problem": "学生信息缺失"})
    station = re.findall(r"收集[\u4e00-\u9fa5]{2,6}站[\u4e00-\u9fa5]{2,12}(?:资料|数据)", text)
    if station:
        errors.append({"type": "content", "problem": "虚构站点采集表述：" + "、".join(station)})
    # 专业/完整性
    if label == "毕业设计成果":
        missing = [r for r in REQUIRED if r not in text]
        if missing:
            errors.append({"type": "professional", "problem": "缺少必备章节关键词：" + "、".join(missing)})
        if "参考文献" not in text:
            errors.append({"type": "professional", "problem": "缺少参考文献"})
        caps = [p for p in doc.paragraphs if re.match(r"^(表|图)\s*\d", p.text.strip())]
        if not caps:
            warnings.append({"type": "professional", "problem": "未检测到表注/图注"})
    # 结构/格式
    st = stats(path)
    if st["max_empty"] > 4:
        warnings.append({"type": "format", "problem": "最大连续空行 %s" % st["max_empty"]})
    if st["breaks"] > 0:
        errors.append({"type": "format", "problem": "人工换行 %s 处" % st["breaks"]})
    return {"file": label, "errors": errors, "warnings": warnings, "stats": st}


def main():
    blueprint = {}
    if os.path.isfile(BLUEPRINT):
        with open(BLUEPRINT, encoding="utf-8") as f:
            blueprint = json.load(f)
    results = [review_file(f) for f in FILES]
    lines = ["# 毕业设计审查报告（周志贤）", "",
             "蓝图章节：%s" % (blueprint.get("章节蓝图") or "未生成"), ""]
    total_errors = 0
    for r in results:
        lines.append("## %s" % r["file"])
        lines.append("- 状态：%s" % ("失败" if r["errors"] else ("通过" if not r["warnings"] else "通过（有警告）")))
        for e in r["errors"]:
            lines.append("- 错误[%s]：%s" % (e["type"], e["problem"]))
        for w in r["warnings"]:
            lines.append("- 警告[%s]：%s" % (w["type"], w["problem"]))
        lines.append("")
        total_errors += len(r["errors"])
    lines.append("## 结论")
    lines.append("")
    lines.append("- 错误总数：%s" % total_errors)
    lines.append("- %s" % ("四文件通过审查" if total_errors == 0 else "存在错误，进入修改-复检循环"))
    report = os.path.join(QUALITY, "毕业设计审查报告_周志贤.md")
    with open(report, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    with open(os.path.join(QUALITY, "毕业设计审查结果_周志贤.json"), "w", encoding="utf-8") as f:
        json.dump(results, f, ensure_ascii=False, indent=2)
    for r in results:
        print(r["file"], "errors", len(r["errors"]), "warnings", len(r["warnings"]))
    print("REVIEWER_DONE", "errors", total_errors)


if __name__ == "__main__":
    main()
